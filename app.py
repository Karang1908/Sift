from __future__ import annotations
import os
import shutil
import json
import asyncio
import ast
import bcrypt
import logging
import re
import secrets
import subprocess
import tempfile
import time
import httpx
from typing import Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Response, Cookie, Depends
from fastapi.responses import StreamingResponse, RedirectResponse, FileResponse
from fastapi.openapi.docs import get_swagger_ui_html, get_redoc_html
from fastapi.openapi.utils import get_openapi
import io
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from parser_utils import extract_text_from_file, is_error_content
import audit_log

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("docparser")

app = FastAPI(title="Sift", docs_url=None, redoc_url=None, openapi_url=None)

# CORS middleware. allow_credentials is deliberately NOT set here: this app
# now issues a session cookie (see Authentication below), and
# allow_origins=["*"] + allow_credentials=True together would let a page on
# ANY other origin ride a logged-in user's cookie and read the response -
# Starlette's CORSMiddleware reflects the real Origin header in that combo
# rather than a literal "*", which is exactly what makes it dangerous once
# there's something worth stealing. The actual frontend is same-origin and
# was never relying on CORS to function, so this is pure hardening.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Paths relative to this file, or Executable directory if packaged.
import sys
def get_writeable_path(name: str, is_file: bool = False) -> str:
    if getattr(sys, 'frozen', False):
        exec_path = sys.executable
        exec_path_lower = exec_path.lower()
        if ".app/contents/macos" in exec_path_lower or "contents/macos" in exec_path_lower:
            # Inside macOS bundle: go up 4 levels to get outside the .app folder
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(exec_path))))
        else:
            # Standalone CLI binary or on Windows: go up 1 level
            base_dir = os.path.dirname(exec_path)
        
        target = os.path.join(base_dir, name)
        
        # Test write permission
        test_dir = base_dir if is_file else target
        try:
            os.makedirs(test_dir, exist_ok=True)
            test_file = os.path.join(test_dir, ".write_test")
            with open(test_file, "w") as f:
                f.write("test")
            os.remove(test_file)
            return target
        except (OSError, PermissionError):
            # Fallback to user home directory under ~/.document_parser/
            home_dir = os.path.expanduser("~")
            fallback_base = os.path.join(home_dir, ".document_parser")
            fallback_target = os.path.join(fallback_base, name)
            fallback_write_dir = fallback_base if is_file else fallback_target
            os.makedirs(fallback_write_dir, exist_ok=True)
            return fallback_target
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        target = os.path.join(base_dir, name)
        if not is_file:
            os.makedirs(target, exist_ok=True)
        return target

if getattr(sys, 'frozen', False):
    BUNDLE_DIR = sys._MEIPASS
    STATIC_DIR = os.path.join(BUNDLE_DIR, "static")
else:
    STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")

# ── Per-user storage ──────────────────────────────────────────────────
#
# Every user's uploads/cache/presets/templates live under their own
# subdirectory (or their own file, for the two preset stores) inside these
# root directories - never a shared pool. Because the directory itself is
# namespaced by username, cross-user access is structurally impossible for
# uploads/cache/templates (there's no shared path to even guess at), not
# just checked after the fact - the simplest correct way to guarantee data
# doesn't leak between accounts. `username` here is always a value that has
# already been validated against USERS by get_current_user() - never raw
# request input - so joining it directly into a path is safe.
UPLOAD_ROOT = get_writeable_path("uploads")
CACHE_ROOT = get_writeable_path("parsed_cache")
PRESETS_ROOT = get_writeable_path("presets")
EXPORT_TEMPLATES_ROOT = get_writeable_path("export_templates")
EXPORT_PRESETS_ROOT = get_writeable_path("export_presets")

# ── Audit trail (admin-only, see audit_log.py) ────────────────────────
#
# Permanent per-user archive - never touched by delete_file() or any
# user-facing action, so it survives whatever the user does in their own
# workspace. Populated at upload/generation time, not at delete time (see
# CLAUDE.md for why). LOGS_ROOT holds the single append-only activity.jsonl
# that backs the admin panel's activity feed.
AUDIT_UPLOADS_ROOT = get_writeable_path("audit_uploads")
AUDIT_ANALYSIS_ROOT = get_writeable_path("audit_analysis")
AUDIT_EXPORTS_ROOT = get_writeable_path("audit_exports")
LOGS_ROOT = get_writeable_path("logs")
ACTIVITY_LOG_PATH = os.path.join(LOGS_ROOT, "activity.jsonl")


def user_upload_dir(username: str) -> str:
    path = os.path.join(UPLOAD_ROOT, username)
    os.makedirs(path, exist_ok=True)
    return path


def user_cache_dir(username: str) -> str:
    path = os.path.join(CACHE_ROOT, username)
    os.makedirs(path, exist_ok=True)
    return path


def user_presets_file(username: str) -> str:
    os.makedirs(PRESETS_ROOT, exist_ok=True)
    return os.path.join(PRESETS_ROOT, f"{username}.json")


# Per-format export instructions (separate preset system from the analysis-query
# presets file above) and the user-uploaded template files they may reference.
# Templates live as random-suffixed files under the user's own subdirectory and
# are unlinked only when the last preset referencing them is deleted - never
# blanket-cleaned (a user might still have a download dialog open).
def user_export_templates_dir(username: str) -> str:
    path = os.path.join(EXPORT_TEMPLATES_ROOT, username)
    os.makedirs(path, exist_ok=True)
    return path


def user_export_presets_file(username: str) -> str:
    os.makedirs(EXPORT_PRESETS_ROOT, exist_ok=True)
    return os.path.join(EXPORT_PRESETS_ROOT, f"{username}.json")


EXPORT_TEMPLATE_FORMATS = {
    ".xlsx": "excel",
    ".docx": "word",
    ".pdf":  "pdf",
}
EXPORT_TEMPLATE_MAX_BYTES = 20 * 1024 * 1024  # 20 MB

OLLAMA_URL = "http://localhost:11434/api/chat"
MODEL_NAME = "minimax-m3:cloud"


# ── Authentication ───────────────────────────────────────────────────
#
# Hardcoded accounts for a small, fixed team - not a public service, so
# there's no signup flow or database, just this dict. To add or change a
# user:
#   1. Generate a bcrypt hash for their password:
#        python3 -c "import bcrypt; print(bcrypt.hashpw(b'their-password', bcrypt.gensalt()).decode())"
#   2. Add/edit an entry below with that hash - NEVER a plaintext password.
#   3. Restart the server (or let --reload pick up the file change).
USERS = {
    "admin": {
        "password_hash": b"$2b$12$lzQFPbHpcC8dJgr2UTKjQuwORJ0ezjXdy2VBrzljnY71xVAuQyl3a", # Password: admin
        "is_admin": True,
    },
    "testuser": {
        "password_hash": b"$2b$12$y8p14vX8fPulwqIf7LZpfO19BjuTn.NS7j/37J5nEoZoGy2ym05YK", # Password: test1234
    },
}

# Used to keep login response time roughly constant whether or not the
# submitted username exists, so failed logins can't be used to enumerate
# valid usernames by timing. Computed once at import time (bcrypt hashing
# is deliberately slow, ~100ms) rather than per-request.
_DUMMY_PASSWORD_HASH = bcrypt.hashpw(b"dummy-password-for-timing", bcrypt.gensalt())

SESSION_COOKIE_NAME = "sift_session"
SESSION_TTL_SECONDS = 6 * 60 * 60  # 6 hours, fixed from login time

# In-memory session store: token -> {"username", "expires_at"}. Lost on
# server restart (including uvicorn --reload) - acceptable for a local/LAN
# tool with 3 known users; a lost session just means logging in again, and
# avoids persisting session tokens to disk.
_sessions: dict = {}

# Brute-force guard: username -> {"failures", "locked_until"}. Not a full
# audit/rate-limiting system - proportionate to "3 known local accounts,"
# not a public-facing service.
LOGIN_MAX_FAILURES = 5
LOGIN_LOCKOUT_SECONDS = 60
_login_attempts: dict = {}


def _prune_sessions_and_lockouts():
    now = time.time()
    expired_sessions = [t for t, s in _sessions.items() if s["expires_at"] < now]
    for t in expired_sessions:
        _sessions.pop(t, None)
    expired_lockouts = [
        u for u, a in _login_attempts.items()
        if a["locked_until"] < now and a.get("updated_at", 0) < now - 3600
    ]
    for u in expired_lockouts:
        _login_attempts.pop(u, None)


def _create_session(username: str) -> str:
    _prune_sessions_and_lockouts()
    token = secrets.token_urlsafe(32)
    _sessions[token] = {"username": username, "expires_at": time.time() + SESSION_TTL_SECONDS}
    return token


def _get_session_user(token: str):
    _prune_sessions_and_lockouts()
    session = _sessions.get(token)
    if not session:
        return None
    if session["expires_at"] < time.time():
        _sessions.pop(token, None)
        return None
    return session["username"]


async def get_current_user(sift_session: str = Cookie(default=None)) -> str:
    if not sift_session:
        raise HTTPException(status_code=401, detail="Not authenticated")
    username = _get_session_user(sift_session)
    if not username:
        raise HTTPException(status_code=401, detail="Session expired or invalid")
    return username


async def get_current_admin(user: str = Depends(get_current_user)) -> str:
    if not USERS.get(user, {}).get("is_admin"):
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


def _validate_admin_target_username(username: str) -> str:
    """Admin routes take `username` from the query/path - unlike every
    other per-user helper in this file, that's request input, not a value
    that already came out of get_current_user(). Restrict it to a real
    configured account before it's ever joined into an audit_log path."""
    if username not in USERS:
        raise HTTPException(status_code=404, detail=f"Unknown user: {username!r}")
    return username

VIDEO_EXTENSIONS = {
    '.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v',
    '.mpg', '.mpeg', '.3gp', '.ts', '.ogv', '.m2ts', '.vob',
}

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.get("/")
def read_root():
    return RedirectResponse(url="/static/index.html")


@app.get("/openapi.json", include_in_schema=False)
async def get_openapi_spec(user: str = Depends(get_current_user)):
    return get_openapi(
        title=app.title,
        version=app.version,
        routes=app.routes,
    )


@app.get("/docs", include_in_schema=False)
async def get_swagger_documentation(user: str = Depends(get_current_user)):
    return get_swagger_ui_html(
        openapi_url="/openapi.json",
        title=app.title + " - Swagger UI",
    )


@app.get("/redoc", include_in_schema=False)
async def get_redoc_documentation(user: str = Depends(get_current_user)):
    return get_redoc_html(
        openapi_url="/openapi.json",
        title=app.title + " - ReDoc",
    )


@app.post("/api/login")
async def login(data: dict, response: Response):
    username = (data.get("username") or "").strip()
    password = data.get("password") or ""

    attempt = _login_attempts.get(username)
    if attempt and attempt["locked_until"] > time.time():
        retry_in = int(attempt["locked_until"] - time.time())
        raise HTTPException(
            status_code=429,
            detail=f"Too many failed attempts. Try again in {retry_in}s.",
        )

    account = USERS.get(username)
    # Always run bcrypt against something, even for an unknown username, so
    # response time doesn't reveal whether the username exists.
    hash_to_check = account["password_hash"] if account else _DUMMY_PASSWORD_HASH
    valid = await asyncio.to_thread(bcrypt.checkpw, password.encode("utf-8"), hash_to_check)
    valid = valid and account is not None

    if not valid:
        attempt = _login_attempts.setdefault(username, {"failures": 0, "locked_until": 0, "updated_at": time.time()})
        attempt["failures"] += 1
        attempt["updated_at"] = time.time()
        if attempt["failures"] >= LOGIN_MAX_FAILURES:
            attempt["locked_until"] = time.time() + LOGIN_LOCKOUT_SECONDS
            attempt["failures"] = 0
        logger.warning("failed login attempt for username=%r", username)
        await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, username, "login_failed")
        raise HTTPException(status_code=401, detail="Invalid username or password")

    _login_attempts.pop(username, None)
    token = _create_session(username)
    response.set_cookie(
        SESSION_COOKIE_NAME, token,
        httponly=True, samesite="lax", max_age=SESSION_TTL_SECONDS, path="/",
    )
    logger.info("login: %s", username)
    await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, username, "login")
    return {"username": username, "is_admin": bool(account.get("is_admin"))}


@app.post("/api/logout")
async def logout(response: Response, sift_session: str = Cookie(default=None)):
    if sift_session:
        username = _get_session_user(sift_session)
        _sessions.pop(sift_session, None)
        if username:
            await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, username, "logout")
    response.delete_cookie(SESSION_COOKIE_NAME, path="/")
    return {"message": "Logged out"}


@app.get("/api/me")
async def me(user: str = Depends(get_current_user)):
    return {"username": user, "is_admin": bool(USERS.get(user, {}).get("is_admin"))}


def _save_upload_sync(file_obj, filepath):
    with open(filepath, "wb") as buffer:
        shutil.copyfileobj(file_obj, buffer)


def _save_template_sync(file_obj, filepath):
    size = 0
    with open(filepath, "wb") as buffer:
        while True:
            chunk = file_obj.read(64 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > EXPORT_TEMPLATE_MAX_BYTES:
                raise ValueError("Template file exceeds maximum allowed size")
            buffer.write(chunk)
        buffer.flush()
        os.fsync(buffer.fileno())
    return size


def _write_cache_atomic(cache_path, text):
    cache_dir = os.path.dirname(cache_path)
    fd, tmp_path = tempfile.mkstemp(dir=cache_dir, suffix=".tmp")
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as f:
            f.write(text)
            f.flush()
            os.fsync(f.fileno())
        os.replace(tmp_path, cache_path)  # atomic on same filesystem
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


def _load_presets_sync(presets_file: str) -> dict:
    if not os.path.exists(presets_file):
        return {}
    try:
        with open(presets_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        # A corrupt/unreadable presets file shouldn't lock the user out of
        # every preset endpoint - degrade to "no presets" instead of 500ing.
        logger.warning("presets file unreadable, treating as empty: %s", e)
        return {}


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    filename = file.filename
    safe_filename = os.path.basename(filename)
    if not safe_filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    _, ext = os.path.splitext(safe_filename.lower())
    if ext in VIDEO_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Video files are not supported ({ext}).")

    filepath = os.path.join(user_upload_dir(user), safe_filename)

    try:
        start = time.perf_counter()
        # Save original file (off the event loop thread)
        await asyncio.to_thread(_save_upload_sync, file.file, filepath)
        save_elapsed = time.perf_counter() - start

        # Parse file text content (off the event loop thread)
        parse_start = time.perf_counter()
        parsed_text = await asyncio.to_thread(extract_text_from_file, filepath)
        parse_elapsed = time.perf_counter() - parse_start

        # Save to parser cache atomically, off the event loop thread
        cache_path = os.path.join(user_cache_dir(user), f"{safe_filename}.txt")
        await asyncio.to_thread(_write_cache_atomic, cache_path, parsed_text)

        size = os.path.getsize(filepath)
        status = "error" if is_error_content(parsed_text) else "parsed"
        logger.info(
            "upload: user=%s %s (%d bytes) save=%.2fs parse=%.2fs status=%s",
            user, safe_filename, size, save_elapsed, parse_elapsed, status,
        )

        # Permanent archive copy - independent of whether the user later
        # deletes this file from their own workspace (see audit_log.py).
        record_id = await asyncio.to_thread(
            audit_log.archive_upload_copy, AUDIT_UPLOADS_ROOT, user, safe_filename, filepath,
        )
        await asyncio.to_thread(
            audit_log.log_activity, ACTIVITY_LOG_PATH, user, "upload",
            filename=safe_filename, size=size, status=status, record_id=record_id,
        )

        return {
            "filename": safe_filename,
            "size": size,
            "status": status
        }
    except Exception as e:
        logger.exception("upload failed: %s", safe_filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        raise HTTPException(status_code=500, detail=f"Failed to process and parse file: {str(e)}")
    finally:
        await file.close()


@app.delete("/api/files/{filename}")
def delete_file(filename: str, user: str = Depends(get_current_user)):
    safe_filename = os.path.basename(filename)
    filepath = os.path.join(user_upload_dir(user), safe_filename)
    cache_path = os.path.join(user_cache_dir(user), f"{safe_filename}.txt")

    deleted = False
    if os.path.exists(filepath):
        os.remove(filepath)
        deleted = True
    if os.path.exists(cache_path):
        os.remove(cache_path)
        deleted = True

    if not deleted:
        raise HTTPException(status_code=404, detail="File not found")

    # Live workspace only - the permanent archive copy from upload time is
    # untouched and stays admin-visible (see audit_log.py).
    audit_log.log_activity(ACTIVITY_LOG_PATH, user, "delete", filename=safe_filename)

    return {"message": f"Successfully deleted {safe_filename}"}


@app.get("/api/files")
def list_files(user: str = Depends(get_current_user)):
    files = []
    upload_dir = user_upload_dir(user)
    cache_dir = user_cache_dir(user)
    if os.path.exists(upload_dir):
        for name in os.listdir(upload_dir):
            filepath = os.path.join(upload_dir, name)
            if os.path.isfile(filepath):
                cache_path = os.path.join(cache_dir, f"{name}.txt")
                if os.path.exists(cache_path):
                    # Only the sentinel prefix matters, not the whole
                    # (possibly huge) cached document text.
                    with open(cache_path, "r", encoding="utf-8") as f:
                        prefix = f.read(200)
                    status = "error" if is_error_content(prefix) else "parsed"
                else:
                    status = "pending"
                files.append({
                    "filename": name,
                    "size": os.path.getsize(filepath),
                    "status": status
                })
    return files


# ── Admin panel ──────────────────────────────────────────────────────
#
# Everything below is gated by get_current_admin (403 for a non-admin
# account) and backs static/admin.html - a standalone page, not wired into
# index.html/script.js. Reads only from the permanent audit_log
# archive/activity log, never from a user's own live workspace - deleting a
# file from a user's own view has zero effect on what shows up here.

@app.get("/api/admin/users")
async def admin_list_users(admin: str = Depends(get_current_admin)):
    results = []
    for name, info in sorted(USERS.items()):
        prompt_file = user_presets_file(name)
        prompt_presets = await asyncio.to_thread(_load_presets_sync, prompt_file)
        
        export_file = user_export_presets_file(name)
        export_presets = await asyncio.to_thread(_load_export_presets_sync, export_file)
        
        formatted_prompts = [
            {"name": k, "prompt": v.get("prompt", ""), "updated_at": v.get("updated_at")}
            for k, v in sorted(prompt_presets.items())
        ]
        formatted_exports = [
            {
                "name": k, 
                "format": v.get("format", ""), 
                "instructions": v.get("instructions", ""),
                "template_filename": v.get("template_filename"),
                "updated_at": v.get("updated_at")
            }
            for k, v in sorted(export_presets.items())
        ]
        
        results.append({
            "username": name,
            "is_admin": bool(info.get("is_admin")),
            "prompt_presets": formatted_prompts,
            "export_presets": formatted_exports
        })
    return results


@app.get("/api/admin/activity")
async def admin_activity(
    username: Optional[str] = None, action: Optional[str] = None, limit: int = 500,
    admin: str = Depends(get_current_admin),
):
    if username is not None:
        _validate_admin_target_username(username)
    return await asyncio.to_thread(
        audit_log.read_activity, ACTIVITY_LOG_PATH, username, action, limit,
    )


@app.get("/api/admin/uploads")
async def admin_list_uploads(username: Optional[str] = None, admin: str = Depends(get_current_admin)):
    if username is not None:
        _validate_admin_target_username(username)
    entries = await asyncio.to_thread(audit_log.list_archive, AUDIT_UPLOADS_ROOT, username)
    # Only the newest archived record for a given (user, filename) pair can
    # ever be "still live" - an older revision superseded by a later upload
    # of the same filename (with or without a delete in between) is never
    # still-live even though a file with that name exists now, since those
    # current bytes belong to the newer record, not this one.
    latest_ts = {}
    for e in entries:
        key = (e["username"], e["filename"])
        latest_ts[key] = max(latest_ts.get(key, 0), e["ts"])
    for e in entries:
        key = (e["username"], e["filename"])
        is_latest = e["ts"] == latest_ts[key]
        e["still_live"] = is_latest and os.path.exists(
            os.path.join(user_upload_dir(e["username"]), e["filename"])
        )
    return entries


@app.get("/api/admin/uploads/{username}/{record_id}/download")
async def admin_download_upload(username: str, record_id: str, admin: str = Depends(get_current_admin)):
    _validate_admin_target_username(username)
    try:
        result = await asyncio.to_thread(
            audit_log.get_archive_file_path, AUDIT_UPLOADS_ROOT, username, record_id,
        )
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid record id")
    if result is None:
        raise HTTPException(status_code=404, detail="Archived upload not found")
    path, original_filename = result
    return FileResponse(path, filename=original_filename)


@app.get("/api/admin/analysis")
async def admin_list_analysis(username: Optional[str] = None, admin: str = Depends(get_current_admin)):
    if username is not None:
        _validate_admin_target_username(username)
    return await asyncio.to_thread(audit_log.list_archive, AUDIT_ANALYSIS_ROOT, username)


@app.get("/api/admin/analysis/{username}/{record_id}")
async def admin_get_analysis(username: str, record_id: str, admin: str = Depends(get_current_admin)):
    _validate_admin_target_username(username)
    try:
        meta = await asyncio.to_thread(audit_log.get_archive_meta, AUDIT_ANALYSIS_ROOT, username, record_id)
        content = await asyncio.to_thread(audit_log.get_analysis_content, AUDIT_ANALYSIS_ROOT, username, record_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid record id")
    if meta is None or content is None:
        raise HTTPException(status_code=404, detail="Archived analysis run not found")
    meta["content"] = content
    return meta


@app.get("/api/admin/exports")
async def admin_list_exports(username: Optional[str] = None, admin: str = Depends(get_current_admin)):
    if username is not None:
        _validate_admin_target_username(username)
    return await asyncio.to_thread(audit_log.list_archive, AUDIT_EXPORTS_ROOT, username)


@app.get("/api/admin/exports/{username}/{record_id}/download")
async def admin_download_export(username: str, record_id: str, admin: str = Depends(get_current_admin)):
    _validate_admin_target_username(username)
    try:
        result = await asyncio.to_thread(
            audit_log.get_archive_file_path, AUDIT_EXPORTS_ROOT, username, record_id,
        )
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid record id")
    if result is None:
        raise HTTPException(status_code=404, detail="Archived export not found")
    path, original_filename = result
    return FileResponse(path, filename=original_filename)


@app.get("/api/admin/export-templates/{username}/{filename}/download")
async def admin_download_export_template(
    username: str, filename: str, original_name: Optional[str] = None, admin: str = Depends(get_current_admin)
):
    _validate_admin_target_username(username)
    if not _is_template_filename_valid(username, filename):
        raise HTTPException(status_code=400, detail="Invalid template filename")
    
    path = _template_path(username, filename)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Template file not found")
        
    download_name = original_name or filename
    return FileResponse(
        path=path,
        filename=download_name,
        media_type="application/octet-stream"
    )


@app.get("/api/presets")
async def list_presets(user: str = Depends(get_current_user)):
    presets = await asyncio.to_thread(_load_presets_sync, user_presets_file(user))
    return [
        {"name": name, "prompt": data["prompt"]}
        for name, data in sorted(presets.items())
    ]


@app.post("/api/presets")
async def save_preset(data: dict, user: str = Depends(get_current_user)):
    name = data.get("name", "").strip()
    prompt = data.get("prompt", "")
    if not name:
        raise HTTPException(status_code=400, detail="Preset name is required")
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Preset prompt is required")

    presets_file = user_presets_file(user)
    presets = await asyncio.to_thread(_load_presets_sync, presets_file)
    presets[name] = {"prompt": prompt, "updated_at": time.time()}
    await asyncio.to_thread(_write_cache_atomic, presets_file, json.dumps(presets, indent=2))
    logger.info("preset saved: user=%s %s (%d chars)", user, name, len(prompt))
    await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, user, "preset_save", name=name)
    return {"name": name, "prompt": prompt}


@app.delete("/api/presets/{name}")
async def delete_preset(name: str, user: str = Depends(get_current_user)):
    presets_file = user_presets_file(user)
    presets = await asyncio.to_thread(_load_presets_sync, presets_file)
    if name not in presets:
        raise HTTPException(status_code=404, detail="Preset not found")
    del presets[name]
    await asyncio.to_thread(_write_cache_atomic, presets_file, json.dumps(presets, indent=2))
    logger.info("preset deleted: user=%s %s", user, name)
    await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, user, "preset_delete", name=name)
    return {"message": f"Preset '{name}' deleted"}


# ── Per-format export-instructions presets (separate from analysis-query presets) ──
#
# A user uploading a custom .xlsx/.docx/.pdf template and saving it as a named
# preset for re-use is the main reason this exists as a separate system - the
# analysis-query presets.json is a flat name->prompt dict, and adding a file
# reference to it would muddy its semantics. The two systems never share state.

def _load_export_presets_sync(export_presets_file: str) -> dict:
    if not os.path.exists(export_presets_file):
        return {}
    try:
        with open(export_presets_file, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        # Same corrupt-file safety as the analysis-query presets: degrade to empty.
        logger.warning("export presets file unreadable, treating as empty: %s", e)
        return {}


def _template_path(username: str, filename: str) -> str:
    """Resolve a template filename to its absolute path inside that user's OWN
    template directory - never any other user's. filename is a bare filename
    produced by the upload endpoint, never user-typed path input, but we still
    defend against path traversal by stripping any directory component. username
    always comes from get_current_user() (never raw request input), so scoping
    to it is what actually prevents one user from reaching another's templates -
    there's no shared directory to escape into in the first place."""
    safe = os.path.basename(filename)
    return os.path.join(user_export_templates_dir(username), safe)


def _is_template_filename_valid(username: str, filename: str) -> bool:
    """True iff filename is a bare, non-traversing name that exists in THIS
    user's templates dir. We never trust the frontend-supplied filename to
    point outside the dir even though the upload endpoint generates it
    server-side, and we never look outside the requesting user's own directory."""
    if not filename or "/" in filename or "\\" in filename or ".." in filename:
        return False
    return os.path.exists(_template_path(username, filename))


def _delete_template_if_orphaned(username: str, filename: str, presets: dict):
    """After deleting a preset, drop its template file too - but only if no
    other preset (belonging to the same user - presets is already that user's
    own dict) still references it. Keeps the templates dir tidy without
    accidentally unlinking a file the user is still using from another preset."""
    if not filename:
        return
    for entry in presets.values():
        if entry.get("template_filename") == filename:
            return  # still in use
    path = _template_path(username, filename)
    if os.path.exists(path):
        try:
            os.remove(path)
            logger.info("template file removed: user=%s %s", username, filename)
        except OSError as e:
            logger.warning("could not remove orphan template %s: %s", filename, e)


def _extract_template_schema(template_path: str, fmt_key: str, max_items: int = 300):
    """Read the uploaded template's ACTUAL existing structure (cell coordinates,
    paragraph/table locations, PDF form fields) and return a compact text
    description tagged with stable [ID=...] locators. Each ID is the exact
    string the field-mapping model must use as a JSON key
    (_generate_field_mapping()), and the exact string the trusted splice
    functions (_splice_xlsx_template() etc.) parse back to find the real
    cell/paragraph/field to edit - so "where does this value go" is decided
    once here and never re-guessed downstream. This mirrors the "extract a
    schema, don't paste raw content" principle from Anthropic's format-clone
    skill: never dump the whole file into the prompt, describe its structure
    compactly and cap it.

    ID schemes (parsed by the matching _splice_*_template function):
    - xlsx:  "<SheetName>!<CellRef>"        e.g. "Report!B4"
    - docx paragraph: "paragraph:<index>"    e.g. "paragraph:3"
    - docx table cell: "table:<t>:<r>:<c>"   e.g. "table:0:2:1"
    - pdf form field:  "field:<name>"        e.g. "field:solar_rate"

    Runs server-side (never inside the untrusted sandbox) since it's reading
    a file the user uploaded through our own upload endpoint, using the same
    trust boundary as parser_utils.py already uses for uploaded documents.

    Returns (schema_text: str, pdf_mode: "form" | "fixed-layout" | None).
    pdf_mode is only meaningful for fmt_key == "pdf".
    """
    if fmt_key == "xlsx":
        import openpyxl as _openpyxl
        wb = _openpyxl.load_workbook(template_path, data_only=True)
        lines = []
        count = 0
        for ws in wb.worksheets:
            lines.append(f"SHEET {ws.title!r} (dimensions={ws.dimensions})")
            if ws.merged_cells.ranges:
                lines.append(f"  merged cells: {', '.join(str(r) for r in ws.merged_cells.ranges)}")
            stop = False
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value in (None, ""):
                        continue
                    if count >= max_items:
                        lines.append(f"  ... (truncated at {max_items} cells; more exist)")
                        stop = True
                        break
                    loc_id = f"{ws.title}!{cell.coordinate}"
                    lines.append(
                        f"  [ID={loc_id}] current={str(cell.value)[:200]!r} "
                        f"(number_format={cell.number_format!r})"
                    )
                    count += 1
                    
                    # Detect and include empty neighbor cells as target slots
                    if count < max_items:
                        try:
                            right_cell = ws.cell(row=cell.row, column=cell.column + 1)
                            if right_cell.value in (None, ""):
                                right_loc_id = f"{ws.title}!{right_cell.coordinate}"
                                lines.append(f"  [ID={right_loc_id}] current='' (blank target cell to the right of {cell.coordinate})")
                                count += 1
                        except Exception:
                            pass
                    
                    if count < max_items:
                        try:
                            below_cell = ws.cell(row=cell.row + 1, column=cell.column)
                            if below_cell.value in (None, ""):
                                below_loc_id = f"{ws.title}!{below_cell.coordinate}"
                                lines.append(f"  [ID={below_loc_id}] current='' (blank target cell below {cell.coordinate})")
                                count += 1
                        except Exception:
                            pass
                if stop:
                    break
            if stop:
                break
        return "\n".join(lines) or "(empty workbook)", None

    if fmt_key == "docx":
        from docx import Document as _Document
        doc = _Document(template_path)
        lines = []
        count = 0
        for i, p in enumerate(doc.paragraphs):
            if not p.text.strip():
                continue
            if count >= max_items:
                lines.append(f"... (truncated at {max_items} items; more exist)")
                return "\n".join(lines), None
            lines.append(f"  [ID=paragraph:{i}] (style={p.style.name!r}) current={p.text[:200]!r}")
            count += 1
        for ti, table in enumerate(doc.tables):
            lines.append(f"TABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    is_empty = not cell.text.strip()
                    if is_empty and ri == 0:
                        continue
                    if count >= max_items:
                        lines.append(f"... (truncated at {max_items} items; more exist)")
                        return "\n".join(lines), None
                    if is_empty:
                        # Find the header text for this column if available
                        header_text = ""
                        try:
                            header_text = table.rows[0].cells[ci].text.strip()
                        except Exception:
                            pass
                        header_desc = f" (target cell under column header {header_text!r})" if header_text else ""
                        lines.append(f"  [ID=table:{ti}:{ri}:{ci}] current=''{header_desc}")
                    else:
                        lines.append(f"  [ID=table:{ti}:{ri}:{ci}] current={cell.text[:200]!r}")
                    count += 1
        return "\n".join(lines) or "(empty document)", None

    if fmt_key == "pdf":
        from pypdf import PdfReader as _PdfReader
        reader = _PdfReader(template_path)
        fields = reader.get_fields()
        lines = []
        if fields:
            lines.append(f"MODE: fillable form ({len(fields)} field(s))")
            for name, f in list(fields.items())[:max_items]:
                lines.append(
                    f"  [ID=field:{name}] type={f.get('/FT')!r}, current={f.get('/V')!r}"
                )
            if len(fields) > max_items:
                lines.append(f"... (truncated at {max_items} fields; more exist)")
            return "\n".join(lines), "form"
        lines.append(f"MODE: fixed-layout, no fillable form fields ({len(reader.pages)} page(s))")
        for i, page in enumerate(reader.pages[:5]):
            text = (page.extract_text() or "")[:1500]
            lines.append(f"  PAGE {i} text:\n{text}")
        if len(reader.pages) > 5:
            lines.append(f"  ... ({len(reader.pages) - 5} more page(s) not shown)")
        return "\n".join(lines), "fixed-layout"

    return "(unknown format)", None


# ── Template clone pipeline: field-mapping + deterministic splice ────────
#
# This is the actual "clone, don't recreate" mechanism (adapted from the
# format-clone skill). For xlsx/docx/pdf-with-form-fields, template exports
# no longer ask the model to write and run a whole Python script that opens
# and edits the file - that leaves too much room for the model to guess
# wrong or rebuild instead of edit. Instead: the model's ONLY job is a
# bounded field-mapping task (given the template's real structure and the
# report, return JSON {location_id: value}), and a small set of trusted,
# backend-authored splice functions (never model-generated, never sandboxed
# - they're regular application code) apply that mapping directly. A
# malformed/incomplete JSON mapping is a normal, recoverable failure (retry
# with the parse error fed back), not a security concern, since the code
# doing the actual file editing is ours, not the model's.
#
# PDF templates with no form fields (fixed-layout) can't be mapped this way
# - there's no discrete field to target, just page text - so those still
# use the older model-writes-a-script overlay path (see _generate_ai_export).

def _extract_json_object(text: str):
    """Pull a JSON object out of a model reply that should contain nothing
    but JSON, but might still have stray code fences or wrapper prose.
    Returns the parsed dict, or None if nothing parseable was found."""
    text = text.strip()
    match = re.search(r"```(?:json)?\s*\n(.*?)```", text, re.DOTALL)
    candidate = match.group(1).strip() if match else text
    try:
        return json.loads(candidate)
    except (json.JSONDecodeError, ValueError):
        pass
    start = candidate.find("{")
    end = candidate.rfind("}")
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(candidate[start:end + 1])
        except (json.JSONDecodeError, ValueError):
            pass
    return None


_FIELD_MAPPING_SYSTEM_PROMPT = """You are a precise field-mapping engine. You are given the exact existing structure of a document template (a list of labeled locations, each tagged with a stable [ID=...] locator) and a Markdown report containing data. Your ONLY job is to map values from the report onto the matching template locations - you do not write any code and you do not generate a document.

For each template location, decide what value (if any) from the report belongs there, based on genuinely matching context (what the location's current content/label suggests it represents) - not just superficially similar names. Even if no explicit user instructions are provided, you must automatically identify logical mappings between the report data and the template structure.

Use these structural and spatial context rules to decide mappings:
1. Spreadsheet cell adjacency: In spreadsheets (IDs containing '!'), the descriptive label is typically in an adjacent cell on the same row. For example, if cell 'Sheet1!A4' contains 'Revenue' and cell 'Sheet1!B4' contains '[VALUE]', '____', or is empty, associate the value for 'Revenue' from the report with the ID 'Sheet1!B4'.
2. Table header context: In tables (IDs starting with 'table:'), the first row (row 0) contains the column headers describing what values should go in subsequent rows (e.g., 'table:0:0:1' is a header, and 'table:0:1:1' is the corresponding data slot).
3. Text placeholders: In text paragraphs, current placeholders (like '[VALUE]', '____', or '{{name}}') are target slots.

How much text to output depends on the location type - read its ID prefix:
- IDs containing "!" (spreadsheet cells) and IDs starting with "field:" (PDF form fields) are ATOMIC value slots - output ONLY the bare data value (e.g. "24.8%", "$1,234.56", "42"), never a label.
- IDs starting with "paragraph:" or "table:" (Word text) often mix a label with an inline placeholder in the SAME run, e.g. current="Solar Conversion Rate: [VALUE]" or current="Solar Conversion Rate: ____". For these, output the COMPLETE replacement text with the placeholder portion substituted and the surrounding label preserved EXACTLY as shown (e.g. "Solar Conversion Rate: 24.8%") - because the caller replaces the location's entire text verbatim, not just a substring. If the current text is ALREADY just a bare placeholder with no label (e.g. current="[VALUE]" or current="____" alone), output just the bare value.

Output ONLY a single JSON object mapping each matched location's exact ID string (copy it exactly as shown, including the "sheet!cell" / "paragraph:N" / "table:t:r:c" / "field:name" form) to its value as a JSON string. No markdown code fences, no explanation, no extra keys, no comments. Example: {"Report!B4": "24.8%", "Report!B5": "84%"}

If truly nothing in the report maps to anything in the template, output an empty JSON object: {}"""


async def _generate_field_mapping(template_schema: str, markdown_text: str, instructions: str,
                                   max_attempts: int = 2, progress_cb=None):
    """Ask the model to map the report's data onto the template's exact
    labeled locations. Returns (mapping_dict_or_None, last_error)."""
    user_content = (
        f"--- TEMPLATE STRUCTURE ---\n{template_schema}\n--- END TEMPLATE STRUCTURE ---\n\n"
    )
    if instructions:
        user_content += f"USER INSTRUCTIONS (follow these when deciding what maps where):\n{instructions}\n\n"
    user_content += f"--- REPORT ---\n{markdown_text}\n--- END REPORT ---"

    messages = [
        {"role": "system", "content": _FIELD_MAPPING_SYSTEM_PROMPT},
        {"role": "user", "content": user_content},
    ]

    last_error = None
    for attempt in range(max_attempts):
        if progress_cb:
            progress_cb(
                "mapping",
                "Mapping report data onto template fields..." if attempt == 0
                else f"Retrying field mapping (attempt {attempt + 1}/{max_attempts})...",
            )
        try:
            async with httpx.AsyncClient(timeout=180.0) as client:
                response = await client.post(
                    OLLAMA_URL,
                    json={"model": MODEL_NAME, "messages": messages, "stream": False},
                )
        except httpx.RequestError as e:
            return None, f"Error connecting to Ollama: {str(e)}"

        if response.status_code != 200:
            return None, f"Ollama error: {response.text}"

        reply = response.json().get("message", {}).get("content", "")
        mapping = _extract_json_object(reply)

        if mapping is None:
            last_error = "Model did not return valid JSON"
        elif not isinstance(mapping, dict):
            last_error = "Model's JSON was not an object"
            mapping = None
        else:
            return mapping, None

        logger.warning("Field mapping attempt %d/%d failed: %s", attempt + 1, max_attempts, last_error)
        messages.append({"role": "assistant", "content": reply})
        messages.append({
            "role": "user",
            "content": f"That wasn't valid JSON ({last_error}). Output ONLY the JSON object, nothing else.",
        })

    return None, last_error


def _splice_xlsx_template(template_path: str, mapping: dict):
    """Apply a {location_id: value} mapping to an existing .xlsx template by
    setting .value on the matched cells only - every other cell, and every
    style/border/fill/merge on the matched cells themselves, is left exactly
    as the template had it. Returns (file_bytes, applied_count)."""
    import openpyxl as _openpyxl
    wb = _openpyxl.load_workbook(template_path)
    applied = 0
    
    # Build lookup table of normalized coordinates to (sheet_name, cell_ref)
    lookup = {}
    for sname in wb.sheetnames:
        ws = wb[sname]
        sname_norm = sname.replace(" ", "").lower()
        for row in ws.iter_rows():
            for cell in row:
                coord = cell.coordinate
                # Fully qualified: e.g. "sheet1!b4"
                lookup[f"{sname_norm}!{coord.lower()}"] = (sname, coord)
                # Spaces in sheet name: "sheet 1!b4"
                lookup[f"{sname.lower()}!{coord.lower()}"] = (sname, coord)
                # Coordinate only: "b4" (useful for single sheet workbooks)
                if coord.lower() not in lookup:
                    lookup[coord.lower()] = (sname, coord)

    for loc_id, value in mapping.items():
        if not loc_id or not isinstance(loc_id, str):
            continue
        norm_key = loc_id.replace(" ", "").lower()
        if norm_key in lookup:
            sname, coord = lookup[norm_key]
            ws = wb[sname]
            coerced, _ = smart_value(str(value))
            ws[coord].value = coerced
            applied += 1
        elif "!" in loc_id:
            parts = loc_id.split("!", 1)
            sheet_part = parts[0].strip().lower()
            cell_part = parts[1].strip().upper()
            matched_sheet = None
            for sname in wb.sheetnames:
                if sname.strip().lower() == sheet_part:
                    matched_sheet = sname
                    break
            if matched_sheet:
                ws = wb[matched_sheet]
                try:
                    ws[cell_part]
                    coerced, _ = smart_value(str(value))
                    ws[cell_part].value = coerced
                    applied += 1
                except Exception:
                    pass
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue(), applied


def _splice_docx_template(template_path: str, mapping: dict):
    """Apply a {location_id: value} mapping to an existing .docx template by
    overwriting text in the FIRST existing run of the matched paragraph/cell
    (preserving that run's font/size/color) and blanking any extra runs in
    the same paragraph so old placeholder text doesn't linger split across
    runs. Returns (file_bytes, applied_count)."""
    from docx import Document as _Document
    doc = _Document(template_path)
    applied = 0

    def _set_paragraph_text(p, text):
        if p.runs:
            p.runs[0].text = text
            for extra in p.runs[1:]:
                extra.text = ""
        else:
            p.add_run(text)

    for loc_id, value in mapping.items():
        if not loc_id or not isinstance(loc_id, str):
            continue
        norm_key = loc_id.replace(" ", "").lower()
        text = str(value)
        if "paragraph" in norm_key:
            num_part = "".join(ch for ch in norm_key if ch.isdigit())
            if num_part:
                try:
                    idx = int(num_part)
                    if idx < len(doc.paragraphs):
                        _set_paragraph_text(doc.paragraphs[idx], text)
                        applied += 1
                except Exception:
                    pass
        elif "table" in norm_key:
            parts = [int(s) for s in re.findall(r"\d+", norm_key)]
            if len(parts) >= 3:
                try:
                    ti, ri, ci = parts[0], parts[1], parts[2]
                    if ti < len(doc.tables):
                        table = doc.tables[ti]
                        if ri < len(table.rows):
                            row = table.rows[ri]
                            if ci < len(row.cells):
                                cell = row.cells[ci]
                                if cell.paragraphs:
                                    _set_paragraph_text(cell.paragraphs[0], text)
                                else:
                                    cell.text = text
                                applied += 1
                except Exception:
                    pass
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue(), applied


def _splice_pdf_form_template(template_path: str, mapping: dict):
    """Apply a {location_id: value} mapping to an existing fillable PDF's
    form fields via pypdf - only field values change, the visual design is
    untouched. Returns (file_bytes, applied_count)."""
    from pypdf import PdfReader as _PdfReader, PdfWriter as _PdfWriter
    reader = _PdfReader(template_path)
    writer = _PdfWriter()
    writer.append(reader)

    fields_lookup = {}
    orig_fields = reader.get_fields() or {}
    for name in orig_fields.keys():
        fields_lookup[name.lower().strip()] = name

    applied = 0
    fields_to_update = {}
    for loc_id, value in mapping.items():
        if not loc_id or not isinstance(loc_id, str):
            continue
        norm_key = loc_id.lower().strip()
        if norm_key.startswith("field:"):
            norm_key = norm_key.split(":", 1)[1].strip()
        
        if norm_key in fields_lookup:
            actual_name = fields_lookup[norm_key]
            fields_to_update[actual_name] = str(value)
            applied += 1

    if fields_to_update:
        for page in writer.pages:
            writer.update_page_form_field_values(page, fields_to_update)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue(), applied


async def _run_template_clone_pipeline(markdown_text: str, fmt: ExportFormat, username: str,
                                        template_filename: str, template_schema: str,
                                        user_instructions: str, progress_cb=None):
    """The real clone pipeline for xlsx/docx/pdf-with-form-fields: map fields
    (one bounded LLM call, one retry), then splice deterministically (no LLM
    involved at all in the actual file edit). Returns (bytes_or_None,
    ai_generated, last_error)."""
    mapping, last_error = await _generate_field_mapping(
        template_schema, markdown_text, user_instructions, progress_cb=progress_cb,
    )
    if mapping is None:
        return None, False, last_error

    if progress_cb:
        progress_cb("splice", f"Filling your {fmt.label} template with the report's data...")

    splice_fn = {
        "xlsx": _splice_xlsx_template,
        "docx": _splice_docx_template,
        "pdf": _splice_pdf_form_template,
    }[fmt.key]

    try:
        file_bytes, applied = await asyncio.to_thread(
            splice_fn, _template_path(username, template_filename), mapping,
        )
    except Exception as e:
        logger.exception("Template splice failed for %s", fmt.key)
        return None, False, f"Template splice failed: {e}"

    if progress_cb:
        progress_cb("done", f"Filled {applied} field(s) in your {fmt.label} template.")
    return file_bytes, True, None


@app.get("/api/export-presets")
async def list_export_presets(user: str = Depends(get_current_user)):
    presets = await asyncio.to_thread(_load_export_presets_sync, user_export_presets_file(user))
    return [
        {
            "name": name,
            "format": data.get("format"),
            "instructions": data.get("instructions", ""),
            "template_filename": data.get("template_filename"),
            "template_original_name": data.get("template_original_name"),
            "updated_at": data.get("updated_at"),
        }
        for name, data in sorted(presets.items())
    ]


@app.post("/api/export-presets")
async def save_export_preset(data: dict, user: str = Depends(get_current_user)):
    name = (data.get("name") or "").strip()
    fmt_key = (data.get("format") or "").strip()
    instructions = data.get("instructions", "")
    template_filename = data.get("template_filename") or None
    template_original_name = data.get("template_original_name") or None

    if not name:
        raise HTTPException(status_code=400, detail="Preset name is required")
    if fmt_key not in ("pdf", "excel", "word"):
        raise HTTPException(status_code=400, detail="format must be 'pdf', 'excel', or 'word'")
    if template_filename and not _is_template_filename_valid(user, template_filename):
        raise HTTPException(status_code=400, detail="template_filename does not exist on the server")

    presets_file = user_export_presets_file(user)
    presets = await asyncio.to_thread(_load_export_presets_sync, presets_file)
    presets[name] = {
        "format": fmt_key,
        "instructions": instructions,
        "template_filename": template_filename,
        "template_original_name": template_original_name,
        "updated_at": time.time(),
    }
    await asyncio.to_thread(
        _write_cache_atomic, presets_file, json.dumps(presets, indent=2),
    )
    logger.info(
        "export preset saved: user=%s %s (format=%s, instructions=%d chars, template=%s)",
        user, name, fmt_key, len(instructions), template_filename or "<none>",
    )
    await asyncio.to_thread(
        audit_log.log_activity, ACTIVITY_LOG_PATH, user, "export_preset_save",
        name=name, format=fmt_key,
    )
    return {
        "name": name,
        "format": fmt_key,
        "instructions": instructions,
        "template_filename": template_filename,
        "template_original_name": template_original_name,
        "updated_at": presets[name]["updated_at"],
    }


@app.delete("/api/export-presets/{name}")
async def delete_export_preset(name: str, user: str = Depends(get_current_user)):
    presets_file = user_export_presets_file(user)
    presets = await asyncio.to_thread(_load_export_presets_sync, presets_file)
    if name not in presets:
        raise HTTPException(status_code=404, detail="Preset not found")
    deleted = presets.pop(name)
    await asyncio.to_thread(
        _write_cache_atomic, presets_file, json.dumps(presets, indent=2),
    )
    # If the deleted preset had a template file, unlink it only when no other
    # preset still references it (see _delete_template_if_orphaned).
    _delete_template_if_orphaned(user, deleted.get("template_filename"), presets)
    logger.info("export preset deleted: user=%s %s", user, name)
    await asyncio.to_thread(audit_log.log_activity, ACTIVITY_LOG_PATH, user, "export_preset_delete", name=name)
    return {"message": f"Export preset '{name}' deleted"}


@app.post("/api/export-templates")
async def upload_export_template(file: UploadFile = File(...), user: str = Depends(get_current_user)):
    """Accept a single .xlsx / .docx / .pdf template file and stash it under
    THIS user's own template directory with a random-suffixed server-side
    filename. The response's `filename` is the value the client should embed
    in a preset's `template_filename` field - never trust a client-supplied
    path, and it only ever resolves inside the uploading user's own directory."""
    original_name = file.filename or ""
    safe_original = os.path.basename(original_name)
    _, ext = os.path.splitext(safe_original.lower())
    if ext not in EXPORT_TEMPLATE_FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Template format {ext!r} is not supported. Use .xlsx, .docx, or .pdf.",
        )

    templates_dir = user_export_templates_dir(user)
    fd, tmp_path = tempfile.mkstemp(dir=templates_dir, suffix=ext)
    os.close(fd)
    try:
        # Read and write the template file on a sync worker thread
        size = await asyncio.to_thread(_save_template_sync, file.file, tmp_path)

        # Move the tmp file to a uuid-suffixed final name so re-uploads don't
        # clobber existing template files referenced by other presets.
        import uuid as _uuid
        final_name = f"{_uuid.uuid4().hex}{ext}"
        final_path = os.path.join(templates_dir, final_name)
        os.replace(tmp_path, final_path)
        logger.info("template uploaded: user=%s %s (saved as %s, %d bytes)", user, safe_original, final_name, size)
        await asyncio.to_thread(
            audit_log.log_activity, ACTIVITY_LOG_PATH, user, "template_upload",
            filename=safe_original, stored_name=final_name, size=size,
        )
        return {
            "filename": final_name,
            "original_name": safe_original,
            "size": size,
            "format": EXPORT_TEMPLATE_FORMATS[ext],
        }
    except ValueError as e:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except OSError: pass
        raise HTTPException(status_code=413, detail=str(e))
    except Exception:
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except OSError: pass
        raise
    finally:
        await file.close()


# ── AI Enhance for export instructions ────────────────────────────────
#
# Mirrors /api/enhance-prompt: a one-shot non-streaming call to Ollama that
# rewrites the user's plain-language export instructions into something the
# downstream export-skill model will follow more precisely. Narrow
# httpx.RequestError catch + status check outside the try are intentional -
# a bare `except Exception` here would catch our own 4xx/5xx HTTPExceptions
# and flatten them to generic 500s (a bug that lived here once).

@app.post("/api/enhance-instructions")
async def enhance_instructions(data: dict, user: str = Depends(get_current_user)):
    instructions = (data.get("instructions") or "").strip()
    fmt_key = (data.get("format") or "").strip()
    if fmt_key not in ("pdf", "excel", "word"):
        raise HTTPException(status_code=400, detail="format must be 'pdf', 'excel', or 'word'")
    if not instructions:
        raise HTTPException(status_code=400, detail="Instructions are required to enhance")

    format_guidance = {
        "pdf": (
            "The enhanced instructions will be passed to a model that writes Python "
            "(reportlab) code to build a polished PDF. Tell the model exactly what "
            "the PDF should look like - page size, margins, color scheme, font, "
            "section structure, whether to use a cover page, how to handle tables "
            "and figures, page numbers, headers/footers, any specific callouts or "
            "highlighting for key figures."
        ),
        "excel": (
            "The enhanced instructions will be passed to a model that writes Python "
            "(openpyxl) code to build a polished Excel workbook. Tell the model "
            "exactly which sheets to create, column order, number formats for any "
            "values that should be percentages/currency/dates, what should be bolded, "
            "frozen panes, autofilter, conditional formatting, named ranges, header "
            "row colors, tab colors, print setup, and any specific aggregation or "
            "pivot-table logic the user wants."
        ),
        "word": (
            "The enhanced instructions will be passed to a model that writes Python "
            "(python-docx) code to build a polished Word document. Tell the model "
            "exactly what the document should contain - title page or not, font and "
            "size, heading style and color scheme, paragraph spacing, bullet vs. "
            "numbered lists, whether to include tables and how they should look, "
            "page headers/footers, page numbering, any logo or letterhead placement, "
            "and the overall document tone (executive brief, formal report, internal "
            "memo, etc.)."
        ),
    }

    system_prompt = (
        "You are an expert prompt engineer who specializes in writing precise, "
        "actionable instructions for AI agents that build documents. The user has "
        "written a short, plain-language request describing how they want their "
        f"document exported. Your job is to rewrite it into a clear, structured "
        f"instruction that a downstream code-writing model can follow exactly.\n\n"
        f"{format_guidance[fmt_key]}\n\n"
        "When enhancing the instructions:\n"
        "1. Restate the user's actual goal in concrete, observable terms (what "
        "the file should LOOK like, not vague goals like 'make it nice').\n"
        "2. Make every value decision explicit - don't say 'appropriate color' "
        "or 'good layout', specify the color hex / table style / font size.\n"
        "3. Preserve the user's original intent and scope - don't add "
        "requirements they didn't ask for, just make what they DID ask for "
        "clearer and more complete.\n"
        "4. If the user said 'use my template' or 'fill in the template', make "
        "explicit that the model must load the existing template file (its path "
        "is passed as a second argument to the build function) and surgically "
        "substitute content into the existing structure rather than building "
        "from scratch.\n"
        "5. The enhanced instructions are plain text the user will read and edit "
        "in a plain textbox - do NOT use Markdown formatting (no #, ##, **, "
        "bullet dashes, pipe tables, code fences). Use plain sentences and "
        "numbered lists only.\n\n"
        "Output ONLY the enhanced instruction text itself. Do not add any "
        "introductory or concluding text (e.g. 'Here is your enhanced instruction...')."
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": instructions},
    ]

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                OLLAMA_URL,
                json={"model": MODEL_NAME, "messages": messages, "stream": False},
            )
    except httpx.RequestError as e:
        raise HTTPException(status_code=500, detail=f"Error connecting to Ollama: {str(e)}")

    if response.status_code != 200:
        raise HTTPException(status_code=response.status_code, detail=f"Ollama error: {response.text}")

    result = response.json()
    enhanced = result.get("message", {}).get("content", "").strip()
    return {"enhanced_instructions": enhanced}


@app.post("/api/enhance-prompt")
async def enhance_prompt(data: dict, user: str = Depends(get_current_user)):
    raw_prompt = data.get("prompt", "")
    if not raw_prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt is required")

    system_prompt = (
        "You are an expert prompt engineer. Your task is to rewrite the user's plain language request "
        "into a highly effective, detailed instruction prompt for a large language model that will "
        "analyze uploaded documents. The analyzing model will have access to the full text content of "
        "multiple files, including page numbers, slide numbers, sheet names, and document names embedded "
        "as citation markers.\n\n"
        "Write the enhanced prompt so it:\n"
        "1. Clearly restates the user's actual goal (e.g., extract specific data, compare/contrast "
        "documents, synthesize findings, summarize) without adding requirements the user didn't ask for.\n"
        "2. Explicitly instructs the analyzing model to cite every claim or quote precisely "
        "(e.g., File Name, Page X, Slide Y, Sheet Z, or Row N).\n"
        "3. If the request involves comparing or contrasting two or more documents, instructs the model "
        "to structure its answer around: an overview of each file, content differences grouped by topic, "
        "key claims, data/evidence and scope/depth, structural or formatting differences between the "
        "files, similarities, a side-by-side quote comparison for major differences (with citations on "
        "both sides), any direct conflicts or contradictions, and a concise summary of the top "
        "differences.\n"
        "4. Tells the model how to organize its final answer for readability (e.g., clear section "
        "headings, bold key terms, comparison tables where useful) — described in plain English "
        "sentences, never by writing literal Markdown syntax.\n"
        "5. Directs the model on how to handle ambiguity, missing data, or conflicting information "
        "across documents (state it explicitly rather than guessing).\n"
        "6. Keeps the original intent and scope of the user's request intact.\n\n"
        "CRITICAL FORMATTING RULE: the enhanced prompt you output is plain text that the user will read "
        "and edit in a plain textbox — it is NOT rendered as Markdown. Never include literal Markdown "
        "syntax in your output (no #, ##, **, bullet dashes, pipe tables, code fences). Describe the "
        "desired structure and emphasis in ordinary prose and simple numbered sentences instead.\n\n"
        "Output ONLY the enhanced prompt text itself. Do not add any introductory or concluding text "
        "(e.g. 'Here is your enhanced prompt...')."
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": raw_prompt}
    ]

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                OLLAMA_URL,
                json={
                    "model": MODEL_NAME,
                    "messages": messages,
                    "stream": False
                }
            )
    except httpx.RequestError as e:
        raise HTTPException(status_code=500, detail=f"Error connecting to Ollama: {str(e)}")

    if response.status_code != 200:
        raise HTTPException(status_code=response.status_code, detail=f"Ollama error: {response.text}")

    result = response.json()
    enhanced = result.get("message", {}).get("content", "").strip()
    return {"enhanced_prompt": enhanced}


@app.post("/api/process")
async def process_files(data: dict, user: str = Depends(get_current_user)):
    prompt = data.get("prompt", "")
    if not prompt.strip():
        raise HTTPException(status_code=400, detail="Prompt is required")

    # Compile parsed text from all of THIS user's uploaded files only - this
    # is the fix for cross-user leakage in the main analysis flow (previously
    # read every .txt in one shared CACHE_DIR, so everyone's documents fed
    # into everyone's analysis).
    cache_dir = user_cache_dir(user)
    document_context = []
    doc_filenames = []
    if os.path.exists(cache_dir):
        for name in sorted(os.listdir(cache_dir)):
            if name.endswith(".txt"):
                filename = name[:-4]  # Remove .txt extension
                filepath = os.path.join(cache_dir, name)
                doc_filenames.append(filename)
                try:
                    with open(filepath, "r", encoding="utf-8") as f:
                        text = f.read()
                    document_context.append(
                        f"--- START OF DOCUMENT: {filename} ---\n"
                        f"{text}\n"
                        f"--- END OF DOCUMENT: {filename} ---"
                    )
                except Exception as e:
                    document_context.append(f"[Error loading document {filename}: {str(e)}]")

    context_str = "\n\n".join(document_context)
    logger.info(
        "process: user=%s prompt_len=%d docs=%d context_chars=%d",
        user, len(prompt), len(document_context), len(context_str),
    )

    system_prompt = (
        "You are an industrial document analyzer. You will receive multiple document texts separated by boundaries.\n"
        "Analyze these documents carefully according to the user's instructions.\n"
        "Provide thorough, precise answers, quoting from the documents and citing specific references (e.g. page, slide, or row numbers) wherever applicable.\n"
        "If the information is not present or cannot be determined, state that clearly."
    )

    user_content = (
        f"Available Documents:\n"
        f"{context_str}\n\n"
        f"User Instructions:\n"
        f"{prompt}"
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content}
    ]

    async def event_generator():
        start = time.perf_counter()
        # Accumulated purely for the permanent audit archive - the generated
        # analysis text was never persisted server-side before (only streamed
        # to and assembled in the browser). Archived in `finally` so a client
        # disconnect mid-stream still saves whatever was generated so far,
        # flagged incomplete rather than silently losing it.
        accumulated = []
        completed_normally = False
        try:
            # read=None: a huge context can legitimately take a long time to
            # ingest before the first token comes back, and generation itself
            # is unbounded by design (no chunking - the whole doc goes in one
            # shot). connect/write stay bounded so a dead Ollama fails fast.
            timeout = httpx.Timeout(connect=10.0, read=None, write=30.0, pool=10.0)
            async with httpx.AsyncClient(timeout=timeout) as client:
                async with client.stream(
                    "POST",
                    OLLAMA_URL,
                    json={
                        "model": MODEL_NAME,
                        "messages": messages,
                        "stream": True
                    }
                ) as response:
                    if response.status_code != 200:
                        yield f"data: {json.dumps({'error': f'Ollama error: status code {response.status_code}'})}\n\n"
                        return

                    async for line in response.aiter_lines():
                        if line:
                            try:
                                chunk = json.loads(line)
                                content = chunk.get("message", {}).get("content", "")
                                done = chunk.get("done", False)
                                accumulated.append(content)
                                if done:
                                    completed_normally = True
                                yield f"data: {json.dumps({'content': content, 'done': done})}\n\n"
                            except Exception as e:
                                yield f"data: {json.dumps({'error': f'Parsing error: {str(e)}'})}\n\n"
            logger.info("process: completed in %.2fs", time.perf_counter() - start)
        except Exception as e:
            logger.exception("process: stream failed after %.2fs", time.perf_counter() - start)
            yield f"data: {json.dumps({'error': f'Stream connection error: {str(e)}'})}\n\n"
        finally:
            full_content = "".join(accumulated)
            if full_content.strip():
                # A client disconnect right around the final chunk cancels
                # THIS task, and asyncio re-raises that cancellation at every
                # subsequent await - including ones inside this finally block.
                # Without shielding, that can abort between the archive write
                # and the log_activity call, leaving an orphaned archive entry
                # with no matching activity-log line (caught live: the .md
                # file existed on disk but "process" never appeared in
                # activity.jsonl). asyncio.shield() keeps the persist task
                # running to completion in the background even if this
                # request's own task is torn down.
                persist_task = asyncio.create_task(_persist_process_archive(
                    user, prompt, full_content, doc_filenames, completed_normally,
                ))
                try:
                    await asyncio.shield(persist_task)
                except asyncio.CancelledError:
                    pass

    return StreamingResponse(event_generator(), media_type="text/event-stream")


async def _persist_process_archive(user: str, prompt: str, full_content: str,
                                     doc_filenames: list, completed_normally: bool) -> None:
    record_id = await asyncio.to_thread(
        audit_log.archive_analysis, AUDIT_ANALYSIS_ROOT, user, prompt,
        full_content, doc_filenames, completed_normally,
    )
    await asyncio.to_thread(
        audit_log.log_activity, ACTIVITY_LOG_PATH, user, "process",
        record_id=record_id, prompt_len=len(prompt), docs=len(doc_filenames),
        complete=completed_normally,
    )


@app.post("/api/export-excel")
async def export_excel(data: dict, user: str = Depends(get_current_user)):
    return await _stream_export(data, _excel_format(), user)


@app.post("/api/export-pdf")
async def export_pdf(data: dict, user: str = Depends(get_current_user)):
    return await _stream_export(data, _pdf_format(), user)


@app.post("/api/export-word")
async def export_word(data: dict, user: str = Depends(get_current_user)):
    return await _stream_export(data, _word_format(), user)


async def _stream_export(data: dict, fmt: ExportFormat, user: str):
    """SSE-streaming export pipeline shared by all formats.

    The frontend opens this with fetch() and gets back a stream of
    `data: {stage, message, pct}\n\n` events as the model is asked,
    the script is validated, executed, and (if needed) the deterministic
    fallback runs. The terminal event carries `{file_b64, filename,
    ai_generated, mime_type}` for the final download.

    Optional request fields:
    - instructions (str): user-provided export instructions. For the
      no-template path, appended to the per-format skill prompt (empty =
      byte-identical to before this feature existed, so panel-3's buttons
      keep working unchanged). For the template clone path, forwarded into
      the field-mapping call instead.
    - template_filename (str | None): bare filename of a template uploaded
      via /api/export-templates. For xlsx, docx, and PDFs with real AcroForm
      fields, this routes to _run_template_clone_pipeline() - the model only
      does a bounded field-mapping task (JSON in, JSON out), and a trusted,
      backend-authored splice function does the actual file edit
      deterministically. Model-written code never touches the template file.
      For fixed-layout PDFs (no form fields - nothing discrete to map onto),
      this falls back to the older _generate_ai_export() script-writing path
      with an explicit lower-fidelity warning. See
      .agents/skills/excel_export/SKILL.md section 6 for the full design.
    """
    markdown_text = data.get("markdown", "")
    if not markdown_text.strip():
        raise HTTPException(status_code=400, detail="Markdown content is required")

    # Stage -> percent mapping drives the progress bar. Trickle between
    # known milestones so the bar always moves even if a stage takes a while
    # (model calls and subprocess execution are the unknown-duration parts).
    stage_pcts = {
        "template": 3, "schema": 8, "mapping": 30, "splice": 70,
        "prompt": 5, "generate": 25, "validate": 55, "execute": 80,
        "fallback": 90, "done": 100,
    }

    user_instructions = (data.get("instructions") or "").strip()
    template_filename = (data.get("template_filename") or "").strip() or None
    if template_filename and not _is_template_filename_valid(user, template_filename):
        raise HTTPException(
            status_code=400,
            detail=f"template_filename {template_filename!r} is not a valid uploaded template",
        )

    template_schema = None
    pdf_mode = None
    template_warning = None
    if template_filename:
        try:
            template_schema, pdf_mode = await asyncio.to_thread(
                _extract_template_schema, _template_path(user, template_filename), fmt.key,
            )
        except Exception as e:
            logger.warning("Could not read template structure for %s: %s", template_filename, e)
            template_schema, pdf_mode = None, None

    # xlsx/docx/pdf-with-form-fields go through the real clone pipeline
    # (field-mapping + deterministic splice). A fixed-layout PDF template
    # has no discrete fields to map onto, so it falls back to the older
    # script-writing overlay path further below.
    use_clone_pipeline = bool(
        template_filename and template_schema is not None
        and (fmt.key in ("xlsx", "docx") or (fmt.key == "pdf" and pdf_mode == "form"))
    )

    # Effective skill prompt for the NON-clone path (no template, or a
    # fixed-layout PDF template). Empty instructions/no template leaves this
    # byte-identical to the format's base prompt, so the no-instructions
    # exports stay unchanged.
    effective_prompt = fmt.skill_prompt
    if user_instructions:
        effective_prompt += (
            "\n\nUSER EXPORT INSTRUCTIONS (follow these in addition to the above):\n"
            + user_instructions
        )
    if template_filename and not use_clone_pipeline and fmt.key == "pdf":
        # Only PDF has a legitimate non-clone template path (fixed-layout
        # overlay). If xlsx/docx land here it's because schema extraction
        # failed - handled separately below, never routed through PDF-only
        # reportlab/pypdf guidance.
        lib_hint = "pypdf.PdfReader(template_path)"
        template_warning = (
            "This PDF template has no fillable form fields, so filling it is lower-fidelity "
            "(best-effort text overlay) than the Excel/Word template paths."
        )
        schema_block = template_schema or "(Could not read template structure)"
        effective_prompt += (
            "\n\nA TEMPLATE FILE WAS UPLOADED for this export. This is a CLONE task, not a "
            "recreate task: the template's EXISTING structure below is the visual contract, and "
            "the report is just data being poured into it.\n\n"
            "IMPORTANT - your build function signature changes for this request: define "
            "build_document(markdown_text: str, template_path: str) (note the second parameter). "
            "The caller will call it with the template's absolute path as that second argument - "
            "do NOT try to find the template via an environment variable, argv, or any file "
            f"search; it is simply the second argument to your function. Open it with {lib_hint}.\n\n"
            f"--- TEMPLATE STRUCTURE ---\n{schema_block}\n--- END TEMPLATE STRUCTURE ---\n\n"
            "This PDF has NO fillable form fields, so true in-place editing isn't possible - "
            "this is a best-effort text overlay, lower fidelity than the Excel/Word template "
            "paths. Use pypdf.PdfReader(template_path) to read the original pages and "
            "reportlab (canvas + a BytesIO buffer) to draw an overlay page with the report's "
            "key values positioned to roughly match where similar text appears in the page "
            "text shown above, then merge the overlay onto the ORIGINAL page with "
            "page.merge_page(overlay_page) - do not discard the original page content, merge "
            "onto it. Write the result with a pypdf.PdfWriter.\n"
            "Map the report's data onto locations that plausibly match the template's page text "
            "shown above - do not add new pages unless the report has more repeated records than "
            "the template shows an example for.\n"
            "Save the modified document to 'output.pdf' as usual."
        )
    elif template_filename and not use_clone_pipeline:
        # xlsx/docx template whose schema extraction failed (corrupt file,
        # unexpected structure, etc.) - there's no safe non-clone fallback
        # for these formats, so proceed template-blind rather than silently
        # misapplying the PDF-only overlay guidance to the wrong format.
        # Surface this to the user instead of pretending the template was used.
        template_warning = (
            f"Could not read your uploaded {fmt.label} template's structure, "
            "so this export was generated without it."
        )

    import base64
    loop = asyncio.get_running_loop()

    async def event_generator():
        queue: asyncio.Queue = asyncio.Queue()

        def progress_cb(stage: str, message: str):
            evt = {
                "stage": stage,
                "message": message,
                "pct": stage_pcts.get(stage, 0),
            }
            try:
                # If running on the same event loop thread, put directly to maintain order
                if asyncio.get_running_loop() == loop:
                    queue.put_nowait(evt)
                    return
            except RuntimeError:
                pass
            loop.call_soon_threadsafe(queue.put_nowait, evt)

        async def run_pipeline():
            try:
                if template_warning:
                    progress_cb("template", template_warning)

                if use_clone_pipeline:
                    file_bytes, ai_generated, last_error = await _run_template_clone_pipeline(
                        markdown_text, fmt, user, template_filename, template_schema,
                        user_instructions, progress_cb=progress_cb,
                    )
                else:
                    file_bytes, ai_generated, last_error = await _generate_ai_export(
                        markdown_text, fmt, progress_cb=progress_cb,
                        instructions=user_instructions,
                        template_filename=template_filename if fmt.key == "pdf" else None,
                        template_owner=user,
                        effective_prompt_override=effective_prompt,
                    )
                if file_bytes is None:
                    logger.warning(
                        "AI %s generation unavailable, using deterministic export: %s",
                        fmt.key, last_error,
                    )
                    progress_cb("fallback", f"AI script failed ({last_error}); using standard formatting...")
                    try:
                        file_bytes = await asyncio.to_thread(fmt.deterministic_fallback, markdown_text)
                        ai_generated = False
                    except Exception as e:
                        logger.exception("Deterministic %s export failed", fmt.key)
                        await queue.put({"stage": "error", "message": f"{fmt.label} generation failed: {e}"})
                        return

                # Permanent archive copy - exported files were never written to
                # disk before (only base64'd into this SSE event and discarded
                # server-side once sent). See audit_log.py. Shielded for the same
                # reason as process_files()'s archive: the outer event_generator
                # explicitly cancels this task on client disconnect, and without
                # shielding that can land between the archive write and the
                # log_activity call, leaving an orphaned archive entry with no
                # matching activity-log line.
                persist_task = asyncio.create_task(_persist_export_archive(
                    user, fmt.key, fmt.download_filename, file_bytes,
                    user_instructions, template_filename, ai_generated,
                ))
                try:
                    await asyncio.shield(persist_task)
                except asyncio.CancelledError:
                    pass

                progress_cb("done", f"{fmt.label} ready.")
                await queue.put({
                    "stage": "file", "filename": fmt.download_filename,
                    "mime_type": fmt.media_type, "ai_generated": ai_generated,
                    "file_b64": base64.b64encode(file_bytes).decode("ascii"),
                })
            except Exception as e:
                logger.exception("Unexpected error in export pipeline task")
                await queue.put({"stage": "error", "message": f"Export pipeline error: {e}"})

        pipeline_task = asyncio.create_task(run_pipeline())

        try:
            while True:
                evt = await queue.get()
                yield f"data: {json.dumps(evt)}\n\n"
                if evt.get("stage") in ("file", "error"):
                    break
        finally:
            if not pipeline_task.done():
                pipeline_task.cancel()
                try:
                    await pipeline_task
                except (asyncio.CancelledError, Exception):
                    pass

    return StreamingResponse(event_generator(), media_type="text/event-stream")


async def _persist_export_archive(user: str, fmt_key: str, filename: str, file_bytes: bytes,
                                    user_instructions: str, template_filename, ai_generated: bool) -> None:
    record_id = await asyncio.to_thread(
        audit_log.archive_export, AUDIT_EXPORTS_ROOT, user, fmt_key,
        filename, file_bytes, user_instructions, template_filename, ai_generated,
    )
    await asyncio.to_thread(
        audit_log.log_activity, ACTIVITY_LOG_PATH, user, "export",
        record_id=record_id, format=fmt_key, ai_generated=ai_generated,
    )


# ── AI-authored Excel generation ─────────────────────────────────────
#
# Instead of forcing every report through one fixed layout, the model reads
# the actual analysis text and writes an openpyxl script tailored to it
# (multi-sheet vs single-sheet, which parts are real tables vs narrative,
# what to format as %/currency/etc). That script is untrusted model output,
# so before it ever runs it is statically restricted to a small safe-import
# allowlist and executed in a separate subprocess (own memory space, no
# filesystem/network access needed by the allowed imports, CPU/memory
# rlimits on POSIX, hard wall-clock timeout) rather than exec()'d in-process.
# This is defense-in-depth, not a real sandbox (no seccomp/chroot/netns) -
# acceptable for a local single-user tool, not something to trust unmodified
# in a multi-tenant context. If generation or validation fails twice, export
# falls back to the deterministic _generate_excel_bytes() below so the
# button never just breaks.

class _ScriptValidationError(Exception):
    pass


# Shared sandbox rules across all AI-generated export scripts. The same
# banned names/attributes block filesystem/process/network dunder escape
# attempts no matter which format is being generated. Per-format differences
# (the specific library the model is allowed to call, the build-function
# name, the output filename) live in the ExportFormat objects below.
_ALLOWED_BASE_IMPORTS = {
    "re", "datetime", "decimal", "math", "string",
    "itertools", "textwrap", "collections", "typing", "io", "dataclasses",
}

_BANNED_NAMES = {
    "os", "sys", "subprocess", "socket", "shutil", "pathlib", "ctypes",
    "importlib", "__import__", "eval", "exec", "compile", "open", "input",
    "exit", "quit", "globals", "locals", "vars", "dir", "breakpoint",
    "requests", "urllib", "http", "ftplib", "smtplib", "multiprocessing",
    "threading", "signal", "resource", "pty", "platform", "getpass",
    "tempfile", "pickle", "marshal", "codecs", "webbrowser", "sqlite3",
}

_BANNED_ATTRS = {
    "__globals__", "__class__", "__bases__", "__subclasses__", "__mro__",
    "__import__", "__builtins__", "__loader__", "__code__", "__closure__",
    "__dict__", "__getattribute__", "__reduce__", "__reduce_ex__",
}


class ExportFormat:
    """One configured export target. All three formats share the same
    AI-writes-code + AST-validate + subprocess-run pipeline, so the only
    per-format things that change are: the skill prompt the model sees, the
    build-function name it must define, the library imports it can use, the
    output filename the wrapper saves to, the media type, and the deterministic
    fallback function for when AI generation isn't available."""
    def __init__(self, key, label, media_type, output_filename, build_fn,
                 allowed_imports, skill_prompt, deterministic_fallback):
        self.key = key
        self.label = label
        self.media_type = media_type
        self.output_filename = output_filename
        self.build_fn = build_fn
        self.allowed_imports = set(allowed_imports) | _ALLOWED_BASE_IMPORTS
        self.skill_prompt = skill_prompt
        self.deterministic_fallback = deterministic_fallback

    @property
    def download_filename(self):
        return f"document_analysis_report.{self.key}"


_EXCEL_SKILL_SYSTEM_PROMPT = """You are an expert data engineer who writes Python (openpyxl) code to turn AI-generated Markdown analysis reports into polished, professional Excel workbooks.

You will be given a Markdown report. Read it carefully and DECIDE the best spreadsheet structure for THIS SPECIFIC content instead of applying one fixed template:
- If the report contains one or more genuinely tabular datasets (markdown tables, comparison grids, structured lists of records), put each into its own sheet as a real Excel table: a proper header row, one row per record, sensible column widths, autofilter, and a frozen header row.
- If the report is mostly narrative analysis (headings, paragraphs, bullet points), lay it out as a readable formatted sheet: merged colored bands for headings, indented markers for bullets, wrapped text for paragraphs.
- If the report compares multiple documents/entities, consider one sheet per entity plus a "Summary" or "Comparison" sheet - whatever best serves someone opening this file in Excel.
- Detect numeric values that were written as plain text and format them correctly: percentages ("45%" -> 0.45 with number_format '0.0%'), currency ("$1,234" -> 1234 with '#,##0.00', right-aligned), integers ('#,##0'), decimals ('#,##0.00'). Never leave numbers as left-aligned plain-text strings.
- Use a clean, professional look: Calibri font, a navy/blue header color scheme, bold white text on colored header bands, subtle alternating row shading in tables, thin borders around table cells, wrapped text for long narrative cells, hidden gridlines, and column widths sized to the content (not the openpyxl default).
- Set print setup on every sheet: portrait orientation, fit-to-width 1 page, 0.5in left/right and 0.6in top/bottom margins.

Write your ENTIRE answer as a single Python code block and nothing else - no explanation before or after.

Hard requirements on the code (violations are rejected and you will be asked to fix and retry):
- Define exactly one top-level function: def build_workbook(markdown_text: str) -> openpyxl.Workbook
- The function must parse `markdown_text` and return a fully populated, styled openpyxl Workbook object. Do not save or write any file yourself - the caller handles saving.
- You may only import from: openpyxl (and submodules, e.g. openpyxl.styles, openpyxl.utils, openpyxl.worksheet.table), re, datetime, decimal, math, string, itertools, textwrap, collections, typing, io, dataclasses. This code runs in an isolated subprocess with no filesystem or network access, so any other import will fail.
- Never use eval, exec, compile, __import__, open, input, or dunder attributes like __class__/__globals__/__subclasses__/__bases__.
- Define any helper functions you need in the same code block, above or inside build_workbook.
- Handle edge cases gracefully (empty sections, no tables at all, malformed markdown) rather than raising - never let build_workbook crash on real input."""


_PDF_SKILL_SYSTEM_PROMPT = """You are an expert document designer who writes Python (reportlab) code to turn AI-generated Markdown analysis reports into polished, professional PDF files.

You will be given a Markdown report. Read it carefully and DECIDE the best document structure for THIS SPECIFIC content instead of applying one fixed template:
- If the report is mostly narrative analysis (headings, paragraphs, bullet points), produce a clean, single-column executive document: a proper title page or styled title, H1/H2/H3 hierarchy, generous paragraph spacing, indented bullet/numbered lists, and page break hints between major sections.
- If the report contains real tabular data (markdown tables, comparison grids), render those as proper reportlab Tables with header rows, alternating row shading, and gridlines - placed inline where the table appears in the markdown, not collected into an appendix.
- If the report compares multiple documents/entities, consider a cover summary section, then per-entity sections with the entities' data laid out clearly, and a side-by-side comparison table.
- If there are key numerical figures (percentages, currencies, counts), make them visually distinct - bold, in a callout box, or in a "Key Figures" table near the top.
- Use a professional, consistent look: a single sans-serif font family (Helvetica or similar), a navy/blue accent color for headings, sensible page margins (~0.75-1in), page numbers in the footer, and a document title in the page header (or just on the first page if the report is short).

Write your ENTIRE answer as a single Python code block and nothing else - no explanation before or after.

Hard requirements on the code (violations are rejected and you will be asked to fix and retry):
- Define exactly one top-level function: def build_document(markdown_text: str) -> reportlab.platypus.BaseDocTemplate or a reportlab.platypus.SimpleDocTemplate instance (already filled with content via its build() call, OR a list of Flowables ready to be passed to a doc's build()).
- The function must parse `markdown_text` and construct a fully populated, styled PDF document. Do not call doc.build() yourself - the caller will read `doc.filename` and call build() on the saved document. Actually, the cleanest contract is: the function returns a SimpleDocTemplate with `filename='output.pdf'` already set, with all Flowables already added to it, but NOT yet built - the caller calls .build(). Alternatively, you can return a list of Flowables and the caller will handle wrapping them in a doc. Pick ONE consistent approach inside your code.
- You may only import from: reportlab (and submodules, e.g. reportlab.platypus, reportlab.lib.styles, reportlab.lib.pagesizes, reportlab.lib.units, reportlab.lib.colors, reportlab.lib.enums, reportlab.platypus.flowables, reportlab.pdfgen), re, datetime, decimal, math, string, itertools, textwrap, collections, typing, io, dataclasses. This code runs in an isolated subprocess with no filesystem or network access, so any other import will fail.
- Never use eval, exec, compile, __import__, open, input, or dunder attributes like __class__/__globals__/__subclasses__/__bases__.
- Define any helper functions you need in the same code block, above or inside build_document.
- Handle edge cases gracefully (empty sections, no tables at all, malformed markdown) rather than raising - never let build_document crash on real input."""


_WORD_SKILL_SYSTEM_PROMPT = """You are an expert document designer who writes Python (python-docx) code to turn AI-generated Markdown analysis reports into polished, professional Microsoft Word (.docx) documents.

You will be given a Markdown report. Read it carefully and DECIDE the best document structure for THIS SPECIFIC content instead of applying one fixed template:
- If the report is mostly narrative analysis (headings, paragraphs, bullet points), produce a clean, single-column executive document: a proper title (large, bold), a styled subtitle/date, an H1/H2/H3 hierarchy, generous paragraph spacing, indented bullet/numbered lists, and a clear page break between major sections.
- If the report contains real tabular data (markdown tables, comparison grids), render those as proper Word tables with a header row, a navy/blue header background with white bold text, alternating row shading, and visible gridlines - placed inline where the table appears in the markdown.
- If the report compares multiple documents/entities, consider a cover summary section, then per-entity sections, and a side-by-side comparison table.
- If there are key numerical figures (percentages, currencies, counts), call them out - in a "Key Figures" table near the top, in bold within the body, or in a styled callout paragraph.
- Use a professional, consistent look: a sans-serif font (Calibri or similar) at 11pt body, navy/blue accent color for headings, sensible page margins (~1in), and add a footer with the page number.

Write your ENTIRE answer as a single Python code block and nothing else - no explanation before or after.

Hard requirements on the code (violations are rejected and you will be asked to fix and retry):
- Define exactly one top-level function: def build_document(markdown_text: str) -> docx.document.Document
- The function must parse `markdown_text` and return a fully populated, styled python-docx Document object. Do not call doc.save() yourself - the caller handles saving.
- You may only import from: docx (and submodules, e.g. docx.shared, docx.enum.text, docx.enum.table, docx.oxml.ns), re, datetime, decimal, math, string, itertools, textwrap, collections, typing, io, dataclasses. This code runs in an isolated subprocess with no filesystem or network access, so any other import will fail.
- Never use eval, exec, compile, __import__, open, input, or dunder attributes like __class__/__globals__/__subclasses__/__bases__.
- Define any helper functions you need in the same code block, above or inside build_document.
- Handle edge cases gracefully (empty sections, no tables at all, malformed markdown) rather than raising - never let build_document crash on real input."""


# Format registry. Fallbacks are added after the deterministic generators
# are defined below, to keep the file's call order forward-referencable.
EXPORT_FORMATS: dict = {}


def _extract_code_block(text: str) -> str:
    match = re.search(r"```(?:python)?\s*\n(.*?)```", text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return text.strip()


def _validate_generated_code(code: str, fmt: ExportFormat) -> None:
    try:
        tree = ast.parse(code)
    except SyntaxError as e:
        raise _ScriptValidationError(f"Python syntax error: {e}")

    if not any(isinstance(node, ast.FunctionDef) and node.name == fmt.build_fn for node in tree.body):
        raise _ScriptValidationError(
            f"Code must define a top-level function named {fmt.build_fn}(markdown_text)."
        )

    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            for alias in node.names:
                root = alias.name.split(".")[0]
                if root not in fmt.allowed_imports:
                    raise _ScriptValidationError(f"Disallowed import: {alias.name}")
        elif isinstance(node, ast.ImportFrom):
            if node.level and node.level > 0:
                raise _ScriptValidationError("Relative imports are not allowed.")
            root = (node.module or "").split(".")[0]
            if root not in fmt.allowed_imports:
                raise _ScriptValidationError(f"Disallowed import: {node.module}")
        elif isinstance(node, ast.Name) and node.id in _BANNED_NAMES:
            raise _ScriptValidationError(f"Disallowed name used: {node.id}")
        elif isinstance(node, ast.Attribute) and node.attr in _BANNED_ATTRS:
            raise _ScriptValidationError(f"Disallowed attribute access: {node.attr}")


def _run_generated_export_script(
    code: str,
    markdown_text: str,
    fmt: ExportFormat,
    username: str,
    template_filename: str | None = None,
) -> bytes:
    """Execute validated, model-authored code in an isolated subprocess.
    The markdown goes in via stdin and the output file comes out at a fixed
    filename in a throwaway cwd - the model's code never needs file/argv
    access at all, which keeps its actual capabilities narrower than what
    the import allowlist alone permits.

    If `template_filename` is set, the template file is copied into the
    sandbox temp dir as 'template.<ext>' and its path is passed as a second
    POSITIONAL ARGUMENT to the model's build function (build_workbook(md, tpl)
    / build_document(md, tpl)) - NOT via an environment variable. os/sys are
    both in _BANNED_NAMES, so a script that tried to read os.environ to find
    the template would always fail validation; passing it as a plain argument
    needs no banned import at all. This was the actual bug behind "the
    template isn't preserved": every template-mode export was silently
    falling back to the deterministic (template-blind) generator because the
    AI path could never pass validation in the first place."""
    scratch_dir = tempfile.mkdtemp(prefix=f"export_{fmt.key}_")
    try:
        # Copy the user template into the sandbox temp dir before running the
        # script. The script only ever sees this controlled path; we never
        # pass an absolute path the user could have tampered with.
        tpl_arg_literal = ""
        if template_filename:
            src = _template_path(username, template_filename)
            dst_name = f"template{os.path.splitext(template_filename)[1]}"
            dst = os.path.join(scratch_dir, dst_name)
            shutil.copyfile(src, dst)
            # repr() so the path is safely embedded as a Python string literal
            # regardless of any special characters in the (backend-controlled,
            # uuid-derived) path.
            tpl_arg_literal = f", {dst!r}"

        if fmt.key == "xlsx":
            invoke = f"_obj = build_workbook(_md{tpl_arg_literal})\n_obj.save('output.xlsx')\n"
        elif fmt.key == "docx":
            invoke = f"_obj = build_document(_md{tpl_arg_literal})\n_obj.save('output.docx')\n"
        elif fmt.key == "pdf":
            # PDF: the build function may return a list of Flowables (no
            # template), a reportlab SimpleDocTemplate (.build()), or a
            # pypdf.PdfWriter (.write() - the form-fill template path uses
            # this one, since PdfWriter has neither .save() nor .build()).
            # Check 'write' before falling back to 'build' so a PdfWriter
            # never mistakenly hits the SimpleDocTemplate branch.
            invoke = (
                f"_obj = build_document(_md{tpl_arg_literal})\n"
                "if isinstance(_obj, list):\n"
                "    from reportlab.platypus import SimpleDocTemplate\n"
                "    _doc = SimpleDocTemplate('output.pdf')\n"
                "    _doc.build(_obj)\n"
                "elif hasattr(_obj, 'write'):\n"
                "    with open('output.pdf', 'wb') as _f:\n"
                "        _obj.write(_f)\n"
                "else:\n"
                "    _obj.filename = 'output.pdf'\n"
                "    _obj.build()\n"
            )
        else:
            raise _ScriptValidationError(f"Unknown format: {fmt.key}")

        wrapper = (
            code
            + "\n\n"
            + "import sys as _sys\n"
            + "_md = _sys.stdin.buffer.read().decode('utf-8', errors='replace')\n"
            + invoke
        )

        script_path = os.path.join(scratch_dir, "generate.py")
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(wrapper)

        run_kwargs = {}
        if os.name == "posix":
            def _limit_resources():
                import resource
                # CPU time only - RLIMIT_AS (virtual address space) was tried
                # and rejected: macOS's dynamic linker/ASLR reserves virtual
                # address space for shared libraries well past 1GB before any
                # real allocation happens, so even `python3 -c "print(1)"`
                # fails to start under a 1GB RLIMIT_AS. The subprocess.run()
                # wall-clock timeout below is the real memory/runaway guard.
                resource.setrlimit(resource.RLIMIT_CPU, (20, 20))
            run_kwargs["preexec_fn"] = _limit_resources

        try:
            result = subprocess.run(
                [sys.executable, script_path],
                input=markdown_text.encode("utf-8"),
                capture_output=True,
                cwd=scratch_dir,
                timeout=30,
                **run_kwargs,
            )
        except subprocess.TimeoutExpired:
            raise _ScriptValidationError("Script timed out after 30 seconds.")

        if result.returncode != 0:
            stderr = result.stderr.decode("utf-8", errors="replace")[-2000:]
            raise _ScriptValidationError(f"Script raised an error:\n{stderr}")

        output_path = os.path.join(scratch_dir, fmt.output_filename)
        if not os.path.exists(output_path):
            raise _ScriptValidationError(f"Script completed but did not produce {fmt.output_filename}.")
        with open(output_path, "rb") as f:
            return f.read()
    finally:
        shutil.rmtree(scratch_dir, ignore_errors=True)


async def _generate_ai_export(markdown_text: str, fmt: ExportFormat, max_attempts: int = 2,
                              progress_cb=None, instructions: str = "",
                              template_filename: str | None = None,
                              template_owner: str | None = None,
                              effective_prompt_override: str | None = None):
    """Ask the model to write a build_*() script for this specific report,
    validate + run it, and retry once with the error fed back if it fails.
    progress_cb (optional) is called with (stage, message) for each
    pipeline step so the SSE endpoint can stream a real progress bar.

    `instructions` and `template_filename` flow through the skill prompt
    builder in _stream_export; the fully-rendered prompt is passed in via
    `effective_prompt_override` so the caller's view of the prompt stays
    authoritative (it shows the same thing the model sees).

    Returns (bytes_or_None, ai_generated, last_error)."""
    if progress_cb:
        progress_cb("prompt", f"Asking the model how to structure your {fmt.label}...")

    # The retry-feedback messages below need to restate the correct function
    # signature - template mode adds a second (template_path) parameter, and
    # reminding the model of the WRONG (1-arg) signature on retry would just
    # reintroduce the bug this fixed (see _run_generated_export_script).
    build_fn_call = f"{fmt.build_fn}(markdown_text, template_path)" if template_filename else f"{fmt.build_fn}(markdown_text)"

    system_prompt = effective_prompt_override if effective_prompt_override is not None else fmt.skill_prompt
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Here is the Markdown analysis report:\n\n{markdown_text}"},
    ]

    last_error = None
    for attempt in range(max_attempts):
        if progress_cb:
            if attempt == 0:
                progress_cb("generate", "Model is writing the export script...")
            else:
                progress_cb("generate", f"Retrying (attempt {attempt + 1}/{max_attempts})...")

        try:
            async with httpx.AsyncClient(timeout=180.0) as client:
                response = await client.post(
                    OLLAMA_URL,
                    json={"model": MODEL_NAME, "messages": messages, "stream": False},
                )
        except httpx.RequestError as e:
            return None, False, f"Error connecting to Ollama: {str(e)}"

        if response.status_code != 200:
            return None, False, f"Ollama error: {response.text}"

        reply = response.json().get("message", {}).get("content", "")
        code = _extract_code_block(reply)

        if progress_cb:
            progress_cb("validate", "Validating the generated script...")

        try:
            await asyncio.to_thread(_validate_generated_code, code, fmt)
        except _ScriptValidationError as e:
            last_error = str(e)
            logger.warning(
                "AI-generated %s script rejected (attempt %d/%d, validation): %s",
                fmt.key, attempt + 1, max_attempts, last_error,
            )
            if progress_cb:
                progress_cb("retry", f"Script rejected: {last_error}")
            messages.append({"role": "assistant", "content": reply})
            messages.append({
                "role": "user",
                "content": (
                    f"That script failed validation: {last_error}\n\n"
                    f"Fix it and output ONLY the corrected Python code block, still defining "
                    f"{build_fn_call} under the same constraints."
                ),
            })
            continue

        if progress_cb:
            progress_cb("execute", "Running the generated script in a sandbox...")

        try:
            file_bytes = await asyncio.to_thread(
                _run_generated_export_script, code, markdown_text, fmt, template_owner, template_filename,
            )
            if progress_cb:
                progress_cb("done", f"AI-authored {fmt.label} ready.")
            return file_bytes, True, None
        except _ScriptValidationError as e:
            last_error = str(e)
            logger.warning(
                "AI-generated %s script rejected (attempt %d/%d, execution): %s",
                fmt.key, attempt + 1, max_attempts, last_error,
            )
            if progress_cb:
                progress_cb("retry", f"Script failed at runtime: {last_error}")
            messages.append({"role": "assistant", "content": reply})
            messages.append({
                "role": "user",
                "content": (
                    f"That script failed at runtime: {last_error}\n\n"
                    f"Fix it and output ONLY the corrected Python code block, still defining "
                    f"{build_fn_call} under the same constraints."
                ),
            })

    return None, False, last_error


def _generate_excel_bytes(markdown_text: str) -> bytes:
    """Convert a Markdown analysis report into a professionally styled single-sheet Excel workbook.
    This is the deterministic fallback used when AI generation isn't available or fails twice -
    it's been working as the only Excel path for a long time and stays correct on its own."""
    import re
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Analysis Report"

    # ── Colour palette ─────────────────────────────────────────────
    NAVY      = "1B2A4A"
    BLUE      = "2E75B6"
    LIGHT_BG  = "F7F9FC"
    ALT_ROW   = "F2F6FC"
    WHITE     = "FFFFFF"
    DARK_TEXT  = "2D3436"
    MID_TEXT   = "636E72"
    BORDER_C  = "D5DDE5"

    # ── Reusable styles ────────────────────────────────────────────
    thin_side    = Side(border_style="thin",   color=BORDER_C)
    medium_side  = Side(border_style="medium", color=BORDER_C)
    cell_border  = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    hdr_border   = Border(left=thin_side, right=thin_side, top=medium_side, bottom=medium_side)

    def clean_md(text: str) -> str:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'__(.+?)__',     r'\1', text)
        text = re.sub(r'_(.+?)_',       r'\1', text)
        text = re.sub(r'`(.+?)`',       r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text.strip()

    def is_table_divider(line: str) -> bool:
        cells = [c.strip() for c in line.strip().split("|")[1:-1]]
        return all(re.match(r'^:?-+:?$', c) or c == '' for c in cells)

    def estimate_row_height(text, span_width_chars=90):
        if not text:
            return 18
        lines_count = 0
        for paragraph in text.split("\n"):
            p_len = len(paragraph)
            lines_count += max(1, (p_len // span_width_chars) + 1)
        return max(18, min(lines_count * 15 + 6, 250))

    # ── First pass: extract tables and narrative blocks ─────────────
    lines = markdown_text.split("\n")
    blocks = []
    i = 0
    max_tbl_cols = 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text  = clean_md(stripped.lstrip("#").strip())
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            header = None
            rows = []
            for tl in table_lines:
                if is_table_divider(tl):
                    continue
                cells = [clean_md(c.strip()) for c in tl.split("|")[1:-1]]
                if header is None:
                    header = cells
                else:
                    rows.append(cells)
            if header:
                max_tbl_cols = max(max_tbl_cols, len(header))
                blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        list_match = re.match(r'^(\s*)([-*+]|\d+[.)]) (.+)', stripped)
        if list_match:
            indent = len(line) - len(line.lstrip())
            marker = list_match.group(2)
            text   = clean_md(list_match.group(3))
            if re.match(r'\d', marker):
                prefix = f"{marker}"
            elif indent >= 2:
                prefix = "  ◦"
            else:
                prefix = "•"
            blocks.append({"type": "bullet", "prefix": prefix, "text": text})
            i += 1
            continue

        if re.match(r'^[-*_]{3,}$', stripped):
            i += 1
            continue

        text = clean_md(stripped)
        if text:
            is_bold = stripped.startswith("**") and stripped.endswith("**")
            blocks.append({"type": "text", "text": text, "bold": is_bold})
        i += 1

    # Standard span width: standard report text merges Columns B to G (span = 6 columns)
    span_cols = max(max_tbl_cols, 6)

    # ── Write Content to Sheet ─────────────────────────────────────
    ws.column_dimensions["A"].width = 5   # narrower column for bullets

    # Initialize columns B, C, D, ... widths
    for col_idx in range(2, 2 + span_cols):
        col_letter = get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = 15  # default, will be adjusted

    r = 2  # start at row 2 for margin padding
    col_max_widths = {}

    def track_col_width(col_idx, val_str):
        col_max_widths[col_idx] = max(col_max_widths.get(col_idx, 0), len(str(val_str or "")))

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            level = block["level"]
            text  = block["text"]

            if r > 2:
                r += 1  # spacer before heading

            # Merge from column A to standard content width
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=1 + span_cols)
            cell = ws.cell(row=r, column=1, value=text)
            cell.alignment = Alignment(horizontal="left", vertical="center")

            if level == 1:
                cell.font = Font(name="Calibri", size=15, bold=True, color=WHITE)
                fill_color = NAVY
                ws.row_dimensions[r].height = 36
            elif level == 2:
                cell.font = Font(name="Calibri", size=13, bold=True, color=WHITE)
                fill_color = BLUE
                ws.row_dimensions[r].height = 30
            else:
                cell.font = Font(name="Calibri", size=11, bold=True, color=NAVY)
                fill_color = LIGHT_BG
                ws.row_dimensions[r].height = 24

            # Apply fill color to merged region
            for c in range(1, 2 + span_cols):
                ws.cell(row=r, column=c).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")

            r += 1

        elif btype == "text":
            text = block["text"]
            # Merge from column B to end of span
            ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=1 + span_cols)
            cell = ws.cell(row=r, column=2, value=text)
            if block.get("bold"):
                cell.font = Font(name="Calibri", size=11, bold=True, color=DARK_TEXT)
            else:
                cell.font = Font(name="Calibri", size=11, color=DARK_TEXT)
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            ws.row_dimensions[r].height = estimate_row_height(text, span_cols * 13)
            r += 1

        elif btype == "bullet":
            text = block["text"]
            prefix = block["prefix"]

            # Bullet marker in Column A
            cell_a = ws.cell(row=r, column=1, value=prefix)
            cell_a.font = Font(name="Calibri", size=11, bold=True, color=BLUE)
            cell_a.alignment = Alignment(horizontal="right", vertical="center")

            # Content merged B to end of span
            ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=1 + span_cols)
            cell_b = ws.cell(row=r, column=2, value=text)
            cell_b.font = Font(name="Calibri", size=11, color=DARK_TEXT)
            cell_b.alignment = Alignment(vertical="center", wrap_text=True)
            ws.row_dimensions[r].height = estimate_row_height(text, span_cols * 13)
            r += 1

        elif btype == "table":
            if r > 2:
                r += 1  # spacer before table

            header = block["header"]
            rows   = block["rows"]
            num_cols = len(header)

            # Write table header starting from Column B
            for ci, h in enumerate(header):
                col_idx = 2 + ci
                cell = ws.cell(row=r, column=col_idx, value=h)
                cell.font      = Font(name="Calibri", size=11, bold=True, color=WHITE)
                cell.fill      = PatternFill(start_color=BLUE, end_color=BLUE, fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border    = hdr_border
                track_col_width(col_idx, h)
            ws.row_dimensions[r].height = 26
            r += 1

            # Write table rows
            for ri, data_row in enumerate(rows):
                is_alt = ri % 2 == 1
                for ci in range(num_cols):
                    col_idx = 2 + ci
                    raw_val = data_row[ci] if ci < len(data_row) else ""
                    val, vtype = smart_value(raw_val)

                    cell = ws.cell(row=r, column=col_idx, value=val)
                    cell.font = Font(name="Calibri", size=11, color=DARK_TEXT)
                    cell.alignment = Alignment(vertical="center", wrap_text=True)
                    cell.border = cell_border
                    if is_alt:
                        cell.fill = PatternFill(start_color=ALT_ROW, end_color=ALT_ROW, fill_type="solid")

                    # Alignment and Formatting
                    if vtype == "pct":
                        cell.number_format = '0.0%'
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    elif vtype == "currency":
                        cell.number_format = '#,##0.00'
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    elif vtype == "int":
                        cell.number_format = '#,##0'
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    elif vtype == "float":
                        cell.number_format = '#,##0.00'
                        cell.alignment = Alignment(horizontal="right", vertical="center")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="center")

                    track_col_width(col_idx, val)

                ws.row_dimensions[r].height = 22
                r += 1

            r += 1  # spacer after table

    # ── Column Width Adjustments ───────────────────────────────────
    for col_idx in range(2, 2 + span_cols):
        col_letter = get_column_letter(col_idx)
        # Find maximum length, clamp between 14 and 45 characters
        length = col_max_widths.get(col_idx, 15)
        ws.column_dimensions[col_letter].width = max(min(length + 4, 45), 14)

    # ── Page setup ─────────────────────────────────────────────────
    ws.sheet_view.showGridLines = False
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.page_margins.top = 0.6
    ws.page_margins.bottom = 0.6

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def smart_value(text: str):
    """Convert a display string to a number/float/percentage/currency where
    possible, mirroring how the value should be stored in a spreadsheet cell
    that already has the matching number_format applied (e.g. "45%" -> 0.45,
    the correct underlying value for a cell formatted as '0.0%'). Shared by
    the deterministic Excel fallback and the template field-splice pipeline
    (_splice_xlsx_template) - a template's cells already carry the right
    number_format, so splicing only needs to get the right underlying value
    into them, never touch the format itself."""
    t = text.strip()
    if not t:
        return t, None
    # Percentage: "45%" → 0.45
    m = re.match(r'^([\d,.]+)\s*%$', t)
    if m:
        try:
            return float(m.group(1).replace(",", "")) / 100.0, "pct"
        except ValueError:
            pass
    # Currency: "$1,234" or "₹500" or "€99.50"
    m = re.match(r'^[$₹€£¥]\s*([\d,.]+)$', t)
    if m:
        try:
            val = float(m.group(1).replace(",", ""))
            return val, "currency"
        except ValueError:
            pass
    # Number with commas: "1,234"
    if re.match(r'^-?[\d,]+$', t) and "," in t:
        try:
            return int(t.replace(",", "")), "int"
        except ValueError:
            pass
    # Integer
    if re.match(r'^-?\d+$', t):
        try:
            return int(t), "int"
        except ValueError:
            pass
    # Float
    if re.match(r'^-?\d+\.\d+$', t):
        try:
            return float(t), "float"
        except ValueError:
            pass
    return t, None


# ── Shared markdown block parser (used by every deterministic fallback) ──
#
# Each format's deterministic generator runs the same block-level pass over
# the markdown first, then renders the resulting list of {"type", ...} dicts
# into its own document model. This is the part where the AI path beats the
# deterministic one in real use: the deterministic path is one-size-fits-all
# (a single sheet, or a single document flow), while the AI path can decide
# to break out a sub-section as its own table/sheet/etc.

def _parse_markdown_blocks(markdown_text: str):
    """Yield block dicts (type=text/heading/bullet/table) from a Markdown
    string. Each block carries a 'level' (headings) or 'rows'/'header' (tables)
    or 'prefix' (bullets). Shared by all deterministic exports so layout
    choices stay consistent across formats."""

    def clean_md(text: str) -> str:
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'__(.+?)__',     r'\1', text)
        text = re.sub(r'_(.+?)_',       r'\1', text)
        text = re.sub(r'`(.+?)`',       r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text.strip()

    def is_table_divider(line: str) -> bool:
        cells = [c.strip() for c in line.strip().split("|")[1:-1]]
        return all(re.match(r'^:?-+:?$', c) or c == '' for c in cells)

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            yield {"type": "heading", "level": level, "text": clean_md(stripped.lstrip("#").strip())}
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            header = None
            rows = []
            for tl in table_lines:
                if is_table_divider(tl):
                    continue
                cells = [clean_md(c.strip()) for c in tl.split("|")[1:-1]]
                if header is None:
                    header = cells
                else:
                    rows.append(cells)
            if header:
                yield {"type": "table", "header": header, "rows": rows}
            continue
        list_match = re.match(r'^(\s*)([-*+]|\d+[.)]) (.+)', stripped)
        if list_match:
            indent = len(line) - len(line.lstrip())
            marker = list_match.group(2)
            text   = clean_md(list_match.group(3))
            if re.match(r'\d', marker):
                prefix = marker
            elif indent >= 2:
                prefix = "  ◦"
            else:
                prefix = "•"
            yield {"type": "bullet", "prefix": prefix, "text": text}
            i += 1
            continue
        if re.match(r'^[-*_]{3,}$', stripped):
            i += 1
            continue
        text = clean_md(stripped)
        if text:
            is_bold = stripped.startswith("**") and stripped.endswith("**")
            yield {"type": "text", "text": text, "bold": is_bold}
        i += 1


# ── Deterministic PDF generator (the fallback) ───────────────────────

def _generate_pdf_bytes(markdown_text: str) -> bytes:
    """Render the same Markdown block parser into a reportlab SimpleDocTemplate.
    A clean, single-section executive document - no per-report structure
    decisions, just reliable layout."""
    try:
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
    except ImportError as e:
        raise RuntimeError(
            f"PDF generation requires reportlab, but it's not installed: {e}"
        )

    NAVY = colors.HexColor("#1B2A4A")
    BLUE = colors.HexColor("#2E75B6")
    ALT  = colors.HexColor("#F2F6FC")

    base = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle("H1", parent=base["Heading1"], textColor=colors.white,
                              backColor=NAVY, fontSize=15, leading=20, spaceBefore=12,
                              spaceAfter=12, leftIndent=8, rightIndent=8, borderPadding=6),
        "h2": ParagraphStyle("H2", parent=base["Heading2"], textColor=colors.white,
                              backColor=BLUE, fontSize=13, leading=18, spaceBefore=10,
                              spaceAfter=10, leftIndent=8, rightIndent=8, borderPadding=5),
        "h3": ParagraphStyle("H3", parent=base["Heading3"], textColor=NAVY,
                              backColor=ALT, fontSize=11, leading=16, spaceBefore=8,
                              spaceAfter=6, leftIndent=6, rightIndent=6, borderPadding=4),
        "body": ParagraphStyle("Body", parent=base["BodyText"], textColor=colors.HexColor("#2D3436"),
                                fontSize=11, leading=15, spaceAfter=6, alignment=TA_LEFT),
        "bullet": ParagraphStyle("Bullet", parent=base["BodyText"], textColor=colors.HexColor("#2D3436"),
                                  fontSize=11, leading=15, leftIndent=14, spaceAfter=3, alignment=TA_LEFT),
    }

    out = io.BytesIO()
    doc = SimpleDocTemplate(
        out, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        title="Document Analysis Report",
    )

    flow = []
    for block in _parse_markdown_blocks(markdown_text):
        btype = block["type"]
        if btype == "heading":
            style_key = f"h{min(block['level'], 3)}"
            flow.append(Paragraph(_xml_escape(block["text"]), styles[style_key]))
        elif btype == "text":
            flow.append(Paragraph(_xml_escape(block["text"]), styles["body"]))
        elif btype == "bullet":
            flow.append(Paragraph(_xml_escape(f"{block['prefix']}  {block['text']}"), styles["bullet"]))
        elif btype == "table":
            data = [block["header"]] + [row + [""] * (len(block["header"]) - len(row)) for row in block["rows"]]
            t = Table(data, hAlign="LEFT", repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), BLUE),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",   (0, 0), (-1, -1), 10),
                ("ALIGN",      (0, 0), (-1, 0), "CENTER"),
                ("VALIGN",     (0, 0), (-1, -1), "MIDDLE"),
                ("GRID",       (0, 0), (-1, -1), 0.4, colors.HexColor("#D5DDE5")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, ALT]),
                ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING",   (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
            ]))
            flow.append(t)
            flow.append(Spacer(1, 0.15 * inch))

    if not flow:
        flow.append(Paragraph(" ", styles["body"]))

    doc.build(flow)
    return out.getvalue()


def _xml_escape(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


# ── Deterministic Word (.docx) generator (the fallback) ──────────────

def _generate_word_bytes(markdown_text: str) -> bytes:
    """Render the Markdown block parser into a single python-docx Document.
    Same one-size-fits-all layout idea as the PDF fallback - a clean
    single-section document with the report's headings, paragraphs, bullets,
    and tables inline."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise RuntimeError(
            f"Word generation requires python-docx, but it's not installed: {e}"
        )

    doc = Document()

    # Page margins + default body font
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    NAVY = RGBColor(0x1B, 0x2A, 0x4A)
    BLUE = RGBColor(0x2E, 0x75, 0xB6)
    DARK = RGBColor(0x2D, 0x34, 0x36)

    for block in _parse_markdown_blocks(markdown_text):
        btype = block["type"]
        if btype == "heading":
            level = block["level"]
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Calibri"
            run.bold = True
            if level == 1:
                run.font.size = Pt(15); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
            elif level == 2:
                run.font.size = Pt(13); run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(5)
            else:
                run.font.size = Pt(11); run.font.color.rgb = NAVY
                p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        elif btype == "text":
            p = doc.add_paragraph()
            run = p.add_run(block["text"])
            run.font.name = "Calibri"; run.font.size = Pt(11); run.font.color.rgb = DARK
            p.paragraph_format.space_after = Pt(6)
        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(block["text"])
            run.font.name = "Calibri"; run.font.size = Pt(11); run.font.color.rgb = DARK
        elif btype == "table":
            num_cols = len(block["header"])
            table = doc.add_table(rows=1 + len(block["rows"]), cols=num_cols)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for ci, h in enumerate(block["header"]):
                hdr[ci].text = ""
                para = hdr[ci].paragraphs[0]
                run = para.add_run(h)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for ri, row in enumerate(block["rows"]):
                cells = table.rows[ri + 1].cells
                for ci in range(num_cols):
                    cells[ci].text = row[ci] if ci < len(row) else ""

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ── Format registry ───────────────────────────────────────────────────
#
# Lazy factory functions (not direct objects) because the deterministic
# fallbacks reference functions defined further down in this file - direct
# reference would be a NameError at module load.

def _excel_format() -> ExportFormat:
    return ExportFormat(
        key="xlsx", label="Excel workbook",
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        output_filename="output.xlsx",
        build_fn="build_workbook",
        allowed_imports={"openpyxl"},
        skill_prompt=_EXCEL_SKILL_SYSTEM_PROMPT,
        deterministic_fallback=_generate_excel_bytes,
    )


def _pdf_format() -> ExportFormat:
    return ExportFormat(
        key="pdf", label="PDF document",
        media_type="application/pdf",
        output_filename="output.pdf",
        build_fn="build_document",
        # pypdf is needed for template-fill mode (reading the existing PDF's
        # form fields / pages and merging an overlay onto them) - without it
        # every template-mode script would be rejected by the AST validator
        # before it could even try to open TEMPLATE_PATH.
        allowed_imports={"reportlab", "pypdf"},
        skill_prompt=_PDF_SKILL_SYSTEM_PROMPT,
        deterministic_fallback=_generate_pdf_bytes,
    )


def _word_format() -> ExportFormat:
    return ExportFormat(
        key="docx", label="Word document",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output_filename="output.docx",
        build_fn="build_document",
        allowed_imports={"docx"},
        skill_prompt=_WORD_SKILL_SYSTEM_PROMPT,
        deterministic_fallback=_generate_word_bytes,
    )

